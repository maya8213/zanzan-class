# -*- coding: utf-8 -*-
from flask import Flask, request, jsonify, send_file, render_template, send_from_directory
from flask_cors import CORS
import os, json, uuid, sqlite3, urllib.parse
from datetime import datetime, timedelta
from calendar import monthrange
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import shutil

app = Flask(__name__, template_folder='templates')
app.secret_key = 'class-manager-secret-key-2026'
CORS(app)

BASE_DIR = r'D:\赞赞可以写入\一句话玩转班级管理'
TEMPLATE_PATH = r'D:\赞赞可以读取\一句话玩转班级管理\班会记录模板.docx'
DB_PATH = os.path.join(BASE_DIR, 'database.db')
UPLOAD_DIR = os.path.join(BASE_DIR, 'uploads')
os.makedirs(BASE_DIR, exist_ok=True)
os.makedirs(UPLOAD_DIR, exist_ok=True)

# 6年数据保留期限（天）
MAX_DATA_DAYS = 6 * 365

def is_workday(date_str):
    """判断是否为法定工作日（周一至周五）"""
    try:
        # 使用中国时区（UTC+8）
        dt = datetime.strptime(date_str, '%Y-%m-%d')
        return dt.weekday() < 5  # 0=周一, 4=周五
    except:
        return False

def get_db():
    conn = sqlite3.connect(DB_PATH)
    conn.row_factory = sqlite3.Row
    return conn

def init_db():
    conn = get_db()
    conn.executescript('''
        CREATE TABLE IF NOT EXISTS students (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            name TEXT NOT NULL UNIQUE,
            gender TEXT,
            score INTEGER DEFAULT 100,
            last_update TEXT
        );
        CREATE TABLE IF NOT EXISTS records (
            id TEXT PRIMARY KEY,
            module TEXT,
            date TEXT,
            data TEXT,
            created TEXT
        );
        CREATE TABLE IF NOT EXISTS photos (
            id TEXT PRIMARY KEY,
            filename TEXT,
            created TEXT
        );
        CREATE TABLE IF NOT EXISTS class_meetings (
            id TEXT PRIMARY KEY,
            theme TEXT,
            content TEXT,
            photos TEXT,
            word_path TEXT,
            created TEXT
        );
        CREATE TABLE IF NOT EXISTS settings (
            key TEXT PRIMARY KEY,
            value TEXT
        );
        CREATE TABLE IF NOT EXISTS dining_defaults (
            student_name TEXT PRIMARY KEY,
            last_status TEXT
        );
    ''')
    conn.commit()
    conn.close()

init_db()

# ---------- 启动时清理超过6年的数据 ----------
def cleanup_old_data():
    cutoff = (datetime.now() - timedelta(days=MAX_DATA_DAYS)).strftime('%Y-%m-%d')
    conn = get_db()
    cur = conn.cursor()
    cur.execute('DELETE FROM records WHERE date < ?', (cutoff,))
    deleted_records = cur.rowcount
    cur.execute('DELETE FROM class_meetings WHERE created < ?', (cutoff + 'T00:00:00',))
    deleted_meetings = cur.rowcount
    conn.commit()
    conn.close()
    if deleted_records or deleted_meetings:
        print(f'🗑️ 清理过期数据：删除了 {deleted_records} 条records，{deleted_meetings} 条班会记录（保留{MAX_DATA_DAYS}天前至今）')

cleanup_old_data()

# ---------- 每周一自动重置分数 ----------
def check_weekly_reset():
    today = datetime.now()
    if today.weekday() != 0:
        return
    year_week = today.strftime('%Y-%W')
    conn = get_db()
    row = conn.execute('SELECT value FROM settings WHERE key = ?', ('last_reset_week',)).fetchone()
    if row is None or row['value'] != year_week:
        conn.execute('UPDATE students SET score = 100')
        conn.execute('INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)', ('last_reset_week', year_week))
        conn.commit()
        print(f'[OK] Weekly score reset: week {year_week}')
    conn.close()

check_weekly_reset()

# ---------- 学生名单 ----------

# ---------- 学情反馈（前置路由，避免被其他模块干扰）----------
@app.route('/api/feedback/generate', methods=['GET'])
def generate_feedback():
    """生成学情反馈Word文档"""
    student_name = request.args.get('student_name', '').strip()
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    if not student_name or not date_from or not date_to:
        return jsonify({'error': '请提供学生姓名、开始日期和结束日期'}), 400

    conn = get_db()
    student_row = conn.execute("SELECT name, gender FROM students WHERE name=?", (student_name,)).fetchone()
    if not student_row:
        conn.close()
        return jsonify({'error': f'学生"{student_name}"不在名单中，请先上传学生名单'}), 400
    student_gender = student_row['gender'] or '男'
    conn.close()

    content = _generate_learning_feedback(student_name, date_from, date_to, student_gender)

    doc = Document()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f'{student_name}的学情反馈报告')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '方正小标宋简体'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f'反馈周期：{date_from} 至 {date_to}')
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(102, 102, 102)
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    p1_title = doc.add_paragraph()
    p1_run = p1_title.add_run('一、考勤与就餐情况')
    p1_run.bold = True
    p1_run.font.size = Pt(12)
    p1_run.font.color.rgb = RGBColor(0, 51, 102)
    p1_content = doc.add_paragraph()
    p1_c = p1_content.add_run(content['paragraph1'])
    p1_c.font.size = Pt(11)

    doc.add_paragraph('')

    p2_title = doc.add_paragraph()
    p2_run = p2_title.add_run('二、课堂表现与家校沟通')
    p2_run.bold = True
    p2_run.font.size = Pt(12)
    p2_run.font.color.rgb = RGBColor(0, 51, 102)
    p2_content = doc.add_paragraph()
    p2_c = p2_content.add_run(content['paragraph2'])
    p2_c.font.size = Pt(11)

    doc.add_paragraph('')

    p3_title = doc.add_paragraph()
    p3_run = p3_title.add_run('三、教育建议')
    p3_run.bold = True
    p3_run.font.size = Pt(12)
    p3_run.font.color.rgb = RGBColor(0, 51, 102)
    p3_content = doc.add_paragraph()
    p3_c = p3_content.add_run(content['paragraph3'])
    p3_c.font.size = Pt(11)

    doc.add_paragraph('')
    doc.add_paragraph('')

    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日")}')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(153, 153, 153)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    safe_name = student_name.replace(' ', '_').replace('/', '_')
    filename = f'学情反馈_{safe_name}_{date_from}_{date_to}.docx'
    out_path = os.path.join(BASE_DIR, filename)
    doc.save(out_path)

    return send_file(out_path,
                     as_attachment=True,
                     download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/')
def index():
    return render_template('index.html')
    return render_template('index.html')

@app.route('/api/students', methods=['GET'])
def get_students():
    conn = get_db()
    rows = conn.execute('SELECT * FROM students ORDER BY id').fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

@app.route('/api/students/upload', methods=['POST'])
def upload_students():
    file = request.files.get('file')
    if not file:
        return jsonify({'error': '没有文件'}), 400
    import io
    from openpyxl import load_workbook
    wb = load_workbook(io.BytesIO(file.read()))
    ws = wb.active
    conn = get_db()
    existing = {r['name']: r['score'] for r in conn.execute('SELECT name, score FROM students').fetchall()}
    conn.execute('DELETE FROM students')
    added = 0
    for row in ws.iter_rows(min_row=2, values_only=True):
        name = row[0]
        gender = str(row[1]).strip() if len(row) > 1 and row[1] else ''
        if name and str(name).strip():
            name = str(name).strip()
            old_score = existing.get(name, 100)
            try:
                conn.execute('INSERT INTO students (name, gender, score, last_update) VALUES (?, ?, ?, ?)',
                           (name, gender, old_score, datetime.now().strftime('%Y-%m-%d')))
                added += 1
            except Exception as e:
                pass
    conn.commit()
    conn.close()
    return jsonify({'added': added})

@app.route('/api/students/template', methods=['GET'])
def download_template():
    from openpyxl import Workbook
    from io import BytesIO
    wb = Workbook()
    ws = wb.active
    ws.title = '学生名单'
    ws.append(['姓名', '性别'])
    ws.append(['示例学生', '男'])
    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
              as_attachment=True, download_name='学生名单模板.xlsx')

# ---------- 考勤/就餐/表现 ----------

@app.route('/api/records/save', methods=['POST'])
def save_records():
    data = request.json
    record_id = str(uuid.uuid4())
    module = data.get('module')
    records_list = data.get('records', [])
    conn = get_db()

    # 保存记录
    conn.execute('''INSERT OR REPLACE INTO records (id, module, date, data, created) VALUES (?, ?, ?, ?, ?)''',
                (record_id, module, data.get('date'),
                 json.dumps(records_list, ensure_ascii=False), datetime.now().isoformat()))

    # 如果是就餐模块，同步保存每个学生的默认状态（仅工作日生效）
    if module == 'dining' and is_workday(data.get('date', '')):
        for r in records_list:
            conn.execute('''INSERT OR REPLACE INTO dining_defaults (student_name, last_status) VALUES (?, ?)''',
                        (r.get('name'), r.get('status', '就餐')))

    conn.commit()
    conn.close()
    return jsonify({'id': record_id})

@app.route('/api/records/get', methods=['GET'])
def get_records():
    module = request.args.get('module')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')
    conn = get_db()
    sql = 'SELECT * FROM records WHERE 1=1'
    params = []
    if module:
        sql += ' AND module = ?'
        params.append(module)
    if date_from:
        sql += ' AND date >= ?'
        params.append(date_from)
    if date_to:
        sql += ' AND date <= ?'
        params.append(date_to)
    sql += ' ORDER BY date DESC'
    rows = conn.execute(sql, params).fetchall()
    conn.close()
    return jsonify([dict(r) for r in rows])

# ---------- 分数更新 ----------

@app.route('/api/scores/update', methods=['POST'])
def update_scores():
    data = request.json
    student_name = data.get('name')
    change = int(data.get('change', 0))
    date_str = data.get('date', datetime.now().strftime('%Y-%m-%d'))
    # 仅在法定工作日更新分数
    if not is_workday(date_str):
        conn = get_db()
        row = conn.execute('SELECT score FROM students WHERE name = ?', (student_name,)).fetchone()
        conn.close()
        return jsonify({'score': row['score'] if row else 100, 'skipped': True})
    conn = get_db()
    row = conn.execute('SELECT score FROM students WHERE name = ?', (student_name,)).fetchone()
    new_score = 100
    if row:
        new_score = max(0, min(150, row[0] + change))
        conn.execute('UPDATE students SET score = ?, last_update = ? WHERE name = ?',
                    (new_score, datetime.now().strftime('%Y-%m-%d'), student_name))
    conn.commit()
    conn.close()
    return jsonify({'score': new_score})

# ---------- 照片上传 ----------

@app.route('/api/photo/upload', methods=['POST'])
def upload_photo():
    file = request.files.get('photo')
    if not file:
        return jsonify({'error': '没有文件'}), 400
    photo_id = str(uuid.uuid4())
    # 从文件名获取扩展名
    orig_name = file.filename or ''
    ext = os.path.splitext(orig_name)[1].lower()
    if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic', '.heif'):
        ext = '.jpg'
    filename = f'{photo_id}{ext}'
    save_path = os.path.join(UPLOAD_DIR, filename)
    file.save(save_path)
    conn = get_db()
    conn.execute('INSERT INTO photos (id, filename, created) VALUES (?, ?, ?)',
                (photo_id, filename, datetime.now().isoformat()))
    conn.commit()
    conn.close()
    return jsonify({'id': filename})  # 返回带扩展名的完整文件名

@app.route('/uploads/<path:filename>')
def uploaded_file(filename):
    return send_from_directory(UPLOAD_DIR, filename)

# ---------- 班会记录 ----------

@app.route('/api/classmeeting/export', methods=['POST'])
def export_meeting():
    data = request.json
    theme = data.get('theme', '')
    content = data.get('content', '')
    photos = data.get('photos', [])
    today_str = datetime.now().strftime("%Y年%m月%d日")

    # AI生成完整的班会内容
    generated = _generate_meeting_content(theme, content)
    final_theme = generated['theme']
    final_content = generated['content']
    final_summary = generated['summary']

    # 生成Word文件
    file_id = str(uuid.uuid4())
    out_path = os.path.join(BASE_DIR, f'{file_id}.docx')

    doc = Document()
    # 标题
    title_para = doc.add_paragraph()
    title_run = title_para.add_run('班 会 记 录')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = '方正小标宋简体'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    # 基本信息表格
    info_items = [
        ('班会主题', final_theme),
        ('时    间', today_str),
        ('地    点', '教室'),
        ('主 持 人', '班主任'),
        ('记 录 人', '班干部'),
    ]
    for label, value in info_items:
        p = doc.add_paragraph()
        label_run = p.add_run(f'{label}：')
        label_run.bold = True
        label_run.font.size = Pt(12)
        val_run = p.add_run(value)
        val_run.font.size = Pt(12)

    doc.add_paragraph('')

    # 会议内容（解析段落，生成格式化的Word文档）
    p_content_title = doc.add_paragraph()
    ct_run = p_content_title.add_run('会议内容：')
    ct_run.bold = True
    ct_run.font.size = Pt(12)

    # 按行解析内容，支持"一、xxx"的章节格式
    lines = final_content.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # 检测是否为一、二、三、四、五级标题
        if line.startswith(('一、', '二、', '三、', '四、', '五、', '六、', '七、', '八、')):
            # 子章节标题
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.bold = True
            r.font.size = Pt(12)
            r.font.color.rgb = RGBColor(0, 51, 102)
        else:
            # 正文内容
            p = doc.add_paragraph()
            r = p.add_run(line)
            r.font.size = Pt(12)

    doc.add_paragraph('')

    # 班会总结
    p_summary_title = doc.add_paragraph()
    st_run = p_summary_title.add_run('班会总结：')
    st_run.bold = True
    st_run.font.size = Pt(12)

    p_summary = doc.add_paragraph()
    sr = p_summary.add_run(final_summary)
    sr.font.size = Pt(12)
    doc.add_paragraph('')
    p_summary = doc.add_paragraph()

    # 插入照片（Word中显示图片占位符）
    if photos:
        doc.add_paragraph('')
        p_photos = doc.add_paragraph()
        p_photos.add_run('现场照片：').bold = True
        for photo_id in photos:
            photo_path = os.path.join(UPLOAD_DIR, photo_id)
            if os.path.exists(photo_path):
                try:
                    doc.add_picture(photo_path, width=Inches(5.5))
                except:
                    doc.add_paragraph(f'（照片：{photo_id}）')
            else:
                doc.add_paragraph(f'（照片：{photo_id}）')

    doc.save(out_path)

    # 写入 records 表
    meeting_id = str(uuid.uuid4())
    today = datetime.now().strftime('%Y-%m-%d')
    conn = get_db()
    conn.execute('''INSERT OR REPLACE INTO records (id, module, date, data, created) VALUES (?, ?, ?, ?, ?)''',
                (meeting_id, 'class_meeting', today,
                 json.dumps({'theme': final_theme, 'content': final_content}, ensure_ascii=False),
                 datetime.now().isoformat()))
    conn.commit()
    conn.close()

    # 生成预览HTML
    download_url = f'/api/meeting/file/{file_id}'
    photo_html = ''
    for pid in photos:
        photo_html += f'<div style="margin:6px"><img src="/uploads/{pid}" style="width:100%;max-width:300px;border-radius:8px;border:1px solid #eee"></div>'
    photo_section = f'<div style="margin-top:16px"><h3 style="font-size:14px;color:#333;margin-bottom:8px">📷 现场照片</h3><div style="display:flex;flex-wrap:wrap;gap:8px">{photo_html if photo_html else "<span style=\"color:#999;font-size:13px\">暂无照片</span>"}</div></div>' if photos or True else ''

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>班会记录预览</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;padding:16px}}
.debug{{background:#fff3cd;padding:12px;border-radius:8px;font-size:12px;color:#856404;margin-bottom:16px;white-space:pre-wrap;word-break:break-all}}
.card{{background:#fff;border-radius:12px;padding:24px;max-width:640px;margin:0 auto;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.title{{text-align:center;font-size:22px;font-weight:700;color:#333;margin-bottom:20px;letter-spacing:4px}}
.row{{display:flex;padding:10px 0;border-bottom:1px solid #eee;font-size:14px}}
.row .label{{font-weight:600;color:#333;min-width:80px}}
.row .value{{color:#666;flex:1;word-break:break-all}}
.sec{{margin-top:16px}}
.sec h3{{font-size:14px;color:#333;margin-bottom:8px;font-weight:600}}
.content-box{{background:#fafafa;padding:14px;border-radius:10px;font-size:13px;color:#444;line-height:2;white-space:pre-wrap;min-height:80px}}
.summary-box{{background:#f6ffed;padding:14px;border-radius:10px;font-size:13px;color:#389e0d;line-height:2}}
.btns{{display:flex;gap:12px;margin-top:20px}}
.btn{{flex:1;padding:14px;text-align:center;border-radius:10px;font-size:15px;text-decoration:none;display:block}}
.btn-download{{background:#1890ff;color:#fff}}
.btn-back{{background:#f0f0f0;color:#666;text-align:center;padding:12px;border-radius:10px;margin-top:12px;display:block;text-decoration:none}}
</style>
</head>
<body>
<div class="debug">【调试信息】收到主题="{theme}" | 收到内容长度={len(content)}字符 | 生成的主题="{final_theme}" | 生成的内容长度={len(final_content)}字符</div>
<div class="card">
  <div class="title">📋 班 会 记 录</div>
  <div class="row"><span class="label">班会主题</span><span class="value">{final_theme}</span></div>
  <div class="row"><span class="label">时　　间</span><span class="value">{today_str}</span></div>
  <div class="row"><span class="label">地　　点</span><span class="value">教室</span></div>
  <div class="row"><span class="label">主 持 人</span><span class="value">班主任</span></div>
  <div class="row"><span class="label">记 录 人</span><span class="value">班干部</span></div>
  <div class="sec">
    <h3>📝 会议内容</h3>
    <div class="content-box">{final_content}</div>
  </div>
  <div class="sec">
    <h3>✅ 班会总结（AI智能生成）</h3>
    <div class="summary-box">{final_summary}</div>
  </div>
  {photo_section}
  <div class="btns">
    <a href="{download_url}" class="btn btn-download">📄 下载Word文件</a>
  </div>
  <a href="javascript:window.close()" class="btn-back">← 返回继续编辑</a>
</div>
</body>
</html>'''

    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}
def _generate_meeting_content(theme, content):
    """根据用户输入的班会主题和内容，AI生成完整的班会记录（主题+内容+总结）
    返回 dict: {theme, content, summary}"""
    # 如果用户已填写主题和内容，以此为基础扩展
    # 否则根据用户提供的任一信息推断
    base_theme = theme.strip() if theme and theme.strip() else None
    base_content = content.strip() if content and content.strip() else None

    if base_theme and base_content:
        # 内容超过30字才视为正文，否则视为关键词，让AI生成完整内容
        if len(base_content) >= 100:
            summary = _make_summary(base_theme, base_content)
            return {'theme': base_theme, 'content': base_content, 'summary': summary}
        else:
            full_content = _make_content(base_theme)
            summary = _make_summary(base_theme, full_content)
            return {'theme': base_theme, 'content': full_content, 'summary': summary}
    elif base_theme:
        # 只填写了主题，扩展内容和总结
        full_content = _make_content(base_theme)
        summary = _make_summary(base_theme, full_content)
        return {'theme': base_theme, 'content': full_content, 'summary': summary}
    elif base_content:
        # 只填写了内容，推断主题并用AI生成完整内容
        inferred_theme = _infer_theme(base_content)
        full_content = _make_content(inferred_theme)
        summary = _make_summary(inferred_theme, full_content)
        return {'theme': inferred_theme, 'content': full_content, 'summary': summary}
    else:
        # 什么都没填，返回默认
        return {
            'theme': '主题班会',
            'content': '本次班会圆满完成，同学们积极参与，认真讨论，收获颇丰。',
            'summary': '本次班会圆满结束，同学们积极参与、热烈讨论，取得了良好效果。希望同学们将班会所学落实到行动中。'
        }

@app.route('/api/meeting/file/<file_id>')
def meeting_download_file(file_id):
    """下载班会记录Word文件"""
    file_path = os.path.join(BASE_DIR, f'{file_id}.docx')
    if not os.path.exists(file_path):
        return '文件不存在或已过期，请重新导出', 404
    return send_file(file_path,
                     as_attachment=True,
                     download_name=f'班会记录_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

def _infer_theme(content):
    """从内容推断班会主题"""
    c = content.lower()
    if any(kw in c for kw in ['安全', '消防', '防溺水', '交通', '火灾', '地震']): return '校园安全教育'
    if any(kw in c for kw in ['学习', '期中', '期末', '考试', '成绩', '学业']): return '学习习惯养成'
    if any(kw in c for kw in ['感恩', '孝敬', '父母', '亲情', '回报']): return '感恩父母'
    if any(kw in c for kw in ['诚信', '诚实', '守信', '作弊']): return '诚信教育'
    if any(kw in c for kw in ['环保', '绿色', '低碳', '节能', '垃圾分类']): return '环保教育'
    if any(kw in c for kw in ['心理', '情绪', '压力', '阳光', '健康']): return '心理健康教育'
    if any(kw in c for kw in ['劳动', '卫生', '清洁', '整理', '家务']): return '劳动教育'
    if any(kw in c for kw in ['体育', '运动', '健康', '锻炼', '跑步']): return '体育与健康'
    if any(kw in c for kw in ['读书', '阅读', '书香', '书籍']): return '读书分享'
    if any(kw in c for kw in ['团结', '友爱', '互助', '集体', '合作']): return '团结互助教育'
    if any(kw in c for kw in ['节日', '中秋', '端午', '春节', '重阳', '清明']): return '传统节日教育'
    if any(kw in c for kw in ['理想', '梦想', '未来', '目标', '志向']): return '理想信念教育'
    if any(kw in c for kw in ['网络', '手机', '游戏', '沉迷']): return '网络安全教育'
    if any(kw in c for kw in ['法律', '守法', '规则', '纪律']): return '法治教育'
    if any(kw in c for kw in ['近视', '眼睛', '视力', '护眼']): return '爱眼护眼'
    return '主题班会'

def _make_content(theme):
    """根据主题生成会议内容（约300字，适合A4纸一页）"""
    t = theme.lower()
    if '安全' in t:
        return '''一、安全意识教育
同学们认真观看了安全教育视频，了解了校园安全的重要性。通过实际案例分析，大家深刻认识到日常生活中潜在的安全隐患。

二、常见安全问题讨论
1. 课间活动安全：不追逐打闹，不做危险游戏
2. 上下学交通安全：遵守交通规则，注意来往车辆
3. 食品安全：拒绝三无食品，养成良好饮食习惯
4. 消防安全：了解灭火器位置，掌握逃生技能

三、自我保护方法
同学们学习了如何在紧急情况下保护自己，包括：遇到危险时如何求助、如何正确拨打报警电话、如何在人群中保护自己等实用知识。

四、承诺与倡议
全体同学共同承诺：从我做起，从小事做起，珍爱生命，注意安全，让安全伴随我们健康成长。'''
    elif '学习' in t or '考试' in t:
        return '''一、学习方法分享
优秀学生代表分享了各自的学习方法和心得体会，包括时间管理、笔记整理、错题复习等实用技巧。

二、学习态度讨论
同学们就"为什么要学习"、"学习的意义是什么"等话题展开热烈讨论。大家认识到，学习不仅是为了考试，更是为了增长知识、提升能力、成为对社会有用的人。

三、制定学习计划
全班同学共同制定了本学期的学习目标，计划通过小组互助、互帮互学的方式，在班级形成良好的学习氛围。

四、表彰与鼓励
对近期学习表现优秀的同学进行了表彰，鼓励其他同学以他们为榜样，勤奋学习，争取进步。

五、师生互动
班主任对同学们提出了殷切期望，希望大家在今后的学习中能够保持积极向上的态度，养成良好的学习习惯。'''
    elif '感恩' in t:
        return '''一、感恩父母
通过视频和图片展示，同学们感受到了父母养育自己的辛苦与不易。很多同学观看时流下了感动的泪水。

二、感恩老师
师生共同回顾了老师辛勤付出的点点滴滴——深夜批改作业、耐心辅导功课、关心学生生活。老师是我们成长道路上的引路人。

三、感恩同学
同学之间的互帮互助、团结友爱也是值得感恩的美好情感。在困难时伸出的援手、在迷茫时给予的鼓励，都是珍贵的友谊。

四、感恩社会
讨论了社会上各行各业人们的付出，让我们拥有平安幸福的生活。引导同学们珍惜当下，心怀感激。

五、感恩行动
同学们制定了感恩行动计划：从现在起，用实际行动表达感恩——帮父母做家务、向老师问好、与同学和睦相处。'''
    elif '诚信' in t:
        return '''一、诚信的含义
通过讲解，同学们理解了诚信就是诚实守信，即说老实话、办老实事、做老实人。诚信是中华民族的传统美德，也是做人的根本。

二、诚信故事分享
同学们积极分享了古今中外的诚信故事，如"曾子杀猪"、"商鞅立木"等，从中感受到了诚信的力量。

三、身边诚信与不诚信现象
大家列举了校园中常见的诚信行为（守时、守信、诚实考试）和不诚信行为（抄袭作业、考试作弊、说谎等），并分析了不诚信行为的危害。

四、诚信宣誓
全体同学庄严宣誓：诚实做人，诚信做事，在今后的学习和生活中坚守诚信底线。

五、诚信承诺书
全班同学在诚信承诺书上签名，承诺从小事做起，做一个诚实守信的好学生。'''
    elif '环保' in t:
        return '''一、认识环境保护
通过图片和视频，展示了当前环境面临的严峻形势——大气污染、水污染、垃圾围城等，让同学们深刻认识到保护环境刻不容缓。

二、低碳生活我能行
同学们讨论了在日常生活中如何践行低碳环保：节约用水用电、减少一次性用品使用、做好垃圾分类、爱护花草树木等。

三、校园环保行动
制定了班级环保公约，号召全体同学从身边小事做起：不乱扔垃圾、爱护公共设施、积极参与校园绿化活动。

四、环保知识竞赛
组织了简单的环保知识小竞赛，在趣味问答中增长了同学们的环保知识。

五、绿色倡议
向全班同学发出绿色生活倡议：让我们共同行动，为建设美丽家园贡献自己的一份力量。'''
    elif '心理' in t or '压力' in t:
        return '''一、认识心理健康
通过讲解，同学们了解了心理健康的重要性——健康不仅包括身体健康，还包括心理健康。心理健康的人能够正确面对压力、保持乐观情绪。

二、识别不良情绪
分析了同学们日常生活中常见的不良情绪（焦虑、抑郁、愤怒、恐惧等），教大家学会识别自己和他人的情绪变化。

三、情绪调节方法
1. 运动疗法：通过体育锻炼释放压力
2. 倾诉疗法：向朋友、家人、老师倾诉
3. 放松疗法：深呼吸、冥想、听音乐
4. 转移疗法：做自己喜欢的事情转移注意力

四、建立支持系统
鼓励同学们主动与家人、朋友、老师沟通，建立良好的人际关系，在遇到困难时能够及时获得支持。

五、阳光心态
倡导同学们以积极乐观的心态面对学习和生活中的挑战，做一个心理健康、内心阳光的新时代少年。'''
    elif '劳动' in t:
        return '''一、劳动的意义
通过讲解，同学们认识到劳动是人类生存和发展的基础，劳动最光荣、劳动最伟大、劳动最美丽。从古至今，无数劳动者用智慧和汗水创造了美好生活。

二、热爱劳动教育
回顾了劳模精神——"宁愿一人脏，换来万人净"的掏粪工人时传祥、"一生修筑三条路"的筑路工人们的故事，激励同学们珍惜劳动成果。

三、我是劳动小能手
同学们分享了自己参与家务劳动和校园劳动的经历，包括打扫教室、整理床铺、洗碗洗衣等力所能及的事情。

四、劳动技能展示
部分同学现场展示了整理书包、叠衣服等基本劳动技能，大家互相学习、共同进步。

五、班级劳动公约
制定了班级劳动公约：每天做一件力所能及的家务，每周参加一次集体劳动，用双手创造美好生活。'''
    elif '团结' in t or '互助' in t:
        return '''一、团结的意义
通过"手指游戏"等互动活动，同学们切身体会到团结的力量——五根手指各有长短，只有握成拳头才能发挥最大作用。

二、团结互助故事
分享了"一个篱笆三个桩，一个好汉三个帮"等团结互助的故事，以及中国女排团队协作夺冠的感人事迹。

三、班级团结现状
分析了班级团结互助方面做得好的地方，以及需要改进的地方。大家一致认为：班级是我家，团结靠大家。

四、团队合作游戏
组织了简单的小组合作游戏，让同学们在实践中体验合作的重要性，学会在团队中发挥自己的作用。

五、团结互助承诺
全班同学共同承诺：在学习和生活中，互帮互助、团结友爱，共同为创建优秀班集体而努力。'''
    elif '读书' in t or '阅读' in t:
        return '''一、为什么要读书
同学们讨论了读书的重要性——"读万卷书，行万里路"，读书能够开阔视野、增长知识、陶冶情操、丰富人生。

二、好书推荐
同学们积极推荐自己喜欢的好书，包括文学名著、科普读物、历史故事等种类，并分享了推荐理由和读后感。

三、读书方法
交流了有效的读书方法：精读与泛读结合、做好读书笔记、写读后感、与他人分享读书心得等。

四、读书习惯
倡导同学们养成每天阅读的好习惯，减少玩手机、看电脑的时间，让阅读成为生活的一部分。

五、班级读书计划
制定了班级读书计划：每月读一本好书，每学期开展一次读书分享会，让书香溢满校园。'''
    elif '理想' in t or '梦想' in t:
        return '''一、放飞梦想
每个人都有自己的梦想，梦想是前进的动力。通过观看航天员、科学家、运动员等不同领域人士实现梦想的故事，激励同学们勇敢追梦。

二、我的梦想
同学们纷纷上台分享自己的梦想——有的想当医生救死扶伤，有的想当老师教书育人，有的想当科学家探索未知……每一个梦想都值得被尊重。

三、梦想与努力
认识到：梦想不是空想，需要通过脚踏实地的努力才能实现。从现在做起，从小事做起，为实现梦想奠定基础。

四、坚持与拼搏
通过"滴水穿石"、"铁杵成针"等故事，告诉同学们实现梦想需要坚持不懈、永不放弃的精神。

五、为梦想加油
同学们写下自己的梦想卡片，贴在班级"梦想墙"上，互相鼓励，共同进步。'''
    elif '网络' in t or '手机' in t or '游戏' in t:
        return '''一、网络安全教育
随着互联网的普及，网络已经成为我们生活的重要组成部分。但网络世界也存在着各种安全隐患，需要我们提高警惕。

二、合理使用网络
讨论了如何正确使用网络：遵守网络道德、不沉溺网络游戏、不浏览不良信息、学会分辨网络信息的真伪。

三、网络成瘾的危害
分析了网络成瘾对青少年身心健康、学业成绩、人际交往等方面的危害，警示同学们引以为戒。

四、制定上网公约
全班同学共同制定了《文明上网公约》：每天上网不超过1小时、不在网络上发表不当言论、保护个人隐私、安全上网。

五、走向户外
倡导同学们多参加户外活动，多与朋友面对面交流，让生活更加丰富多彩。'''
    else:
        return f'''一、班会背景
本次{theme}主题班会，旨在引导同学们深入了解{theme}的相关知识，培养良好品格和行为习惯。

二、活动过程
1. 班主任通过图片、视频等多媒体形式，向同学们介绍了{theme}的重要性和意义
2. 同学们认真听讲，积极思考，踊跃发言
3. 小组讨论环节，大家热烈交流，分享各自的理解和心得
4. 各小组代表发言，总结讨论成果

三、收获与感悟
通过本次班会，同学们对{theme}有了更加深刻的认识。大家纷纷表示，要将学到的知识落实到日常行动中，从身边小事做起，争做文明好学生。

四、总结与展望
班主任对本次班会进行了总结，肯定了同学们的积极参与和精彩表现。希望同学们以此次班会为契机，在今后的学习和生活中，不断进步、健康成长。'''

def _make_summary(theme, content):
    """根据主题和内容生成班会总结"""
    t = theme.lower()
    summaries = {
        '安全': '本次安全主题班会显著增强了同学们的安全意识。通过案例学习和讨论，大家深刻认识到安全无小事，要时刻绷紧安全这根弦。希望同学们将安全意识内化于心、外化于行，在日常生活中时刻注意自身和他人的安全，让安全伴随健康成长。',
        '学习': '本次学习主题班会有效激发了同学们的学习热情。通过方法分享和目标制定，大家对学习的意义有了更深刻的理解，对如何提高学习效率也有了更清晰的认识。相信同学们会以此次班会为新起点，勤奋学习、努力拼搏，在学业上取得更大进步。',
        '感恩': '本次感恩主题班会触动了同学们的心灵，让大家深刻体会到感恩的重要。通过分享和讨论，同学们学会了表达感恩、传递温暖。希望同学们常怀感恩之心，用实际行动回报父母、老师、同学和社会的关爱。',
        '诚信': '本次诚信主题班会强化了同学们的诚信意识。通过故事分享和问题讨论，大家认识到诚信是立身之本、处世之道。全体同学庄严承诺：将坚守诚信底线，做诚实守信的好学生，让诚信之花在校园绽放。',
        '环保': '本次环保主题班会增强了同学们的生态文明意识。通过学习和讨论，大家认识到保护环境是每个人的责任。希望同学们从身边小事做起，践行绿色低碳生活方式，为建设美丽中国贡献力量。',
        '心理': '本次心理健康主题班会帮助同学们掌握了情绪调节的方法。通过学习和互动，大家学会了以积极心态面对生活中的压力和挑战。希望同学们保持阳光心态，健康快乐地成长。',
        '劳动': '本次劳动主题班会培养了同学们热爱劳动的观念。通过技能展示和公约制定，大家认识到劳动最光荣、劳动最伟大。希望同学们积极参与劳动，在劳动中锻炼自己、收获成长。',
        '团结': '本次团结互助主题班会增强了班级凝聚力。通过游戏和讨论，同学们深刻体会到团结的力量。希望同学们在今后的学习和生活中互帮互助、携手进步，共同创造美好的班集体。',
        '读书': '本次读书分享主题班会激发了同学们的阅读兴趣。通过好书推荐和方法交流，大家收获了宝贵的读书经验。希望同学们养成良好的阅读习惯，让书香浸润心灵、点亮人生。',
        '理想': '本次理想信念主题班会让同学们坚定了人生方向。通过梦想分享和故事感悟，大家认识到梦想需要坚持和努力才能实现。希望同学们志存高远、脚踏实地，为实现梦想而不懈奋斗。',
        '网络': '本次网络安全教育主题班会让同学们认识到文明上网的重要性。通过学习和讨论，大家学会了如何正确使用网络、远离网络危害。希望同学们合理安排上网时间，让网络成为学习和生活的好帮手。',
    }
    for key, summary in summaries.items():
        if key in t:
            return summary
    return f'本次{theme}主题班会圆满成功，同学们积极参与、热烈讨论。通过本次班会，大家对{theme}有了更深入的认识和理解，增强了相关意识和能力，达到了预期效果。希望同学们将班会所学落实到行动中，不断进步、健康成长。'

@app.route('/api/download', methods=['GET'])
def download():
    path = request.args.get('path')
    if path and os.path.exists(path):
        return send_file(path, as_attachment=True)
    return jsonify({'error': '文件不存在'}), 404

# ---------- 分数重置（每周一）----------

@app.route('/api/scores/reset', methods=['POST'])
def reset_scores():
    conn = get_db()
    conn.execute('UPDATE students SET score = 100')
    year_week = datetime.now().strftime('%Y-%W')
    conn.execute('INSERT OR REPLACE INTO settings (key, value) VALUES (?, ?)', ('last_reset_week', year_week))
    conn.commit()
    conn.close()
    return jsonify({'ok': True, 'week': year_week})

@app.route('/api/scores/status', methods=['GET'])
def scores_status():
    conn = get_db()
    row = conn.execute("SELECT value FROM settings WHERE key = 'last_reset_week'").fetchone()
    last_reset = row['value'] if row else None
    current_week = datetime.now().strftime('%Y-%W')
    is_monday = datetime.now().weekday() == 0
    conn.close()
    return jsonify({
        'last_reset_week': last_reset,
        'current_week': current_week,
        'needs_reset': is_monday and last_reset != current_week
    })

# ---------- 就餐模块：获取本月用餐次数 + 上次默认状态 ----------
@app.route('/api/dining/stats', methods=['GET'])
def dining_stats():
    """获取当前月份每个学生的用餐次数统计"""
    today = datetime.now()
    year = today.year
    month = today.month
    _, last_day = monthrange(year, month)
    date_from = f'{year}-{month:02d}-01'
    date_to = f'{year}-{month:02d}-{last_day:02d}'

    conn = get_db()
    # 获取本月所有就餐记录
    rows = conn.execute(
        'SELECT date, data FROM records WHERE module = ? AND date >= ? AND date <= ?',
        ('dining', date_from, date_to)
    ).fetchall()

    counts = {}  # name -> count（仅统计工作日）
    for row in rows:
        # 仅统计法定工作日（周一至周五）
        if not is_workday(row['date']):
            continue
        try:
            records_list = json.loads(row['data'])
            for r in records_list:
                name = r.get('name')
                status = r.get('status', '就餐')
                if name and status == '就餐':
                    counts[name] = counts.get(name, 0) + 1
        except:
            pass

    # 获取每个学生上次保存的默认状态
    defaults = {}
    rows2 = conn.execute('SELECT student_name, last_status FROM dining_defaults').fetchall()
    for row in rows2:
        defaults[row['student_name']] = row['last_status']

    conn.close()
    return jsonify({'counts': counts, 'defaults': defaults, 'month': f'{year}-{month:02d}'})

# ---------- 考勤模块：获取本月出勤统计 ----------
@app.route('/api/attendance/stats', methods=['GET'])
def attendance_stats():
    """获取当前月份每个学生的出勤统计（病假/事假天数）"""
    today = datetime.now()
    year = today.year
    month = today.month
    _, last_day = monthrange(year, month)
    date_from = f'{year}-{month:02d}-01'
    date_to = f'{year}-{month:02d}-{last_day:02d}'

    conn = get_db()
    rows = conn.execute(
        'SELECT date, data FROM records WHERE module = ? AND date >= ? AND date <= ?',
        ('attendance', date_from, date_to)
    ).fetchall()

    sick_leave = {}   # name -> cumulative sick leave days
    personal_leave = {}  # name -> cumulative personal leave days

    for row in rows:
        if not is_workday(row['date']):
            continue
        try:
            records_list = json.loads(row['data'])
            for r in records_list:
                name = r.get('name')
                status = r.get('status', '正常')
                if not name:
                    continue
                if status in ('病假0.5日', '病假1日'):
                    days = 0.5 if '0.5' in status else 1
                    sick_leave[name] = sick_leave.get(name, 0) + days
                elif status in ('事假0.5日', '事假1日'):
                    days = 0.5 if '0.5' in status else 1
                    personal_leave[name] = personal_leave.get(name, 0) + days
        except:
            pass

    conn.close()
    return jsonify({
        'sick_leave': sick_leave,
        'personal_leave': personal_leave,
        'month': f'{year}-{month:02d}'
    })

# ---------- 表现模块：获取本周得分 ----------
def get_week_range():
    """获取本周一和周日的日期"""
    today = datetime.now()
    monday = today - timedelta(days=today.weekday())
    sunday = monday + timedelta(days=6)
    return monday.strftime('%Y-%m-%d'), sunday.strftime('%Y-%m-%d')

@app.route('/api/performance/weekly', methods=['GET'])
def performance_weekly():
    """获取本周每个学生的表现得分（每周100分，表扬+1，批评-1）"""
    date_from, date_to = get_week_range()
    conn = get_db()
    rows = conn.execute(
        'SELECT date, data FROM records WHERE module = ? AND date >= ? AND date <= ?',
        ('performance', date_from, date_to)
    ).fetchall()
    conn.close()

    scores = {}  # name -> net score change from records (+1 praise, -1 criticism)
    for row in rows:
        try:
            records_list = json.loads(row['data'])
            for r in records_list:
                name = r.get('name', '')
                status = r.get('status', '')
                if not name:
                    continue
                if name not in scores:
                    scores[name] = 0
                if status == '表扬':
                    scores[name] += 1
                elif status == '批评':
                    scores[name] -= 1
        except:
            pass

    return jsonify({'scores': scores, 'week_from': date_from, 'week_to': date_to, 'base': 100})

# ---------- 数据导出（按模块 + 日期区间）----------

@app.route('/api/export/module', methods=['GET'])
def export_module():
    """按模块和日期区间导出数据"""
    module = request.args.get('module', '')  # attendance, dining, performance
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    if not date_from or not date_to:
        return jsonify({'error': '请提供 date_from 和 date_to 参数'}), 400

    conn = get_db()
    sql = 'SELECT * FROM records WHERE 1=1'
    params = []
    if module:
        sql += ' AND module = ?'
        params.append(module)
    sql += ' AND date >= ? AND date <= ?'
    params.extend([date_from, date_to])
    sql += ' ORDER BY date DESC'
    rows = conn.execute(sql, params).fetchall()
    conn.close()

    module_names = {'attendance': '考勤', 'dining': '就餐', 'performance': '表现'}
    return jsonify({
        'module': module,
        'module_name': module_names.get(module, module),
        'date_from': date_from,
        'date_to': date_to,
        'records': [dict(r) for r in rows],
        'export_time': datetime.now().isoformat(),
        'total': len(rows)
    })

@app.route('/api/export/excel', methods=['GET'])
def export_excel():
    """导出为Excel文件（考勤：透视表格式；就餐/表现：列表格式）"""
    module = request.args.get('module', '')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    if not date_from or not date_to:
        return jsonify({'error': '请提供日期区间'}), 400

    conn = get_db()
    sql = 'SELECT * FROM records WHERE 1=1'
    params = []
    if module:
        sql += ' AND module = ?'
        params.append(module)
    sql += ' AND date >= ? AND date <= ?'
    params.extend([date_from, date_to])
    sql += ' ORDER BY date ASC'
    rows = conn.execute(sql, params).fetchall()
    # 获取正式学生名单，导出时只包含名单内的学生
    valid_students = set(row['name'] for row in conn.execute('SELECT name FROM students'))
    conn.close()

    module_names = {'attendance': '考勤', 'dining': '就餐', 'performance': '表现'}
    module_name = module_names.get(module, module)

    from openpyxl import Workbook
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.title = module_name

    if module == 'attendance':
        # ---------- 考勤：透视表格式 ----------
        # 收集所有学生和日期
        all_students = {}  # name -> {date -> status}
        all_dates = []      # 按时间排序的唯一日期

        for row in rows:
            try:
                records_list = json.loads(row['data'])
                d = row['date']
                if d not in all_dates:
                    all_dates.append(d)
                for r in records_list:
                    name = r.get('name', '').strip()
                    if not name or name not in valid_students:
                        continue
                    if name not in all_students:
                        all_students[name] = {}
                    all_students[name][d] = r.get('status', '正常')
            except:
                pass

        all_dates.sort()

        # 表头：姓名 | 日期1 | 日期2 | ... | 汇总
        header = ['姓名'] + all_dates + ['汇总']
        ws.append(header)

        # 数据行
        for name in sorted(all_students.keys()):
            date_statuses = all_students[name]
            # 汇总统计
            sick_days = personal_days = late_count = normal_count = 0
            for d in all_dates:
                status = date_statuses.get(d, '正常')
                if status == '正常':
                    normal_count += 1
                elif status in ('病假0.5日', '病假1日'):
                    sick_days += 0.5 if '0.5' in status else 1
                elif status in ('事假0.5日', '事假1日'):
                    personal_days += 0.5 if '0.5' in status else 1
                elif status == '迟到':
                    late_count += 1

            summary_parts = []
            if sick_days > 0:
                summary_parts.append(f'病假{sick_days}日')
            if personal_days > 0:
                summary_parts.append(f'事假{personal_days}日')
            if late_count > 0:
                summary_parts.append(f'迟到{late_count}次')
            if normal_count == len(all_dates) and not summary_parts:
                summary_parts = ['全勤']
            summary = ' '.join(summary_parts) if summary_parts else ''

            row_data = [name] + [date_statuses.get(d, '') for d in all_dates] + [summary]
            ws.append(row_data)

    elif module == 'dining':
        # ---------- 就餐：透视表格式 ----------
        all_students = {}
        all_dates = []

        for row in rows:
            try:
                records_list = json.loads(row['data'])
                d = row['date']
                if d not in all_dates:
                    all_dates.append(d)
                for r in records_list:
                    name = r.get('name', '').strip()
                    if not name or name not in valid_students:
                        continue
                    if name not in all_students:
                        all_students[name] = {}
                    all_students[name][d] = r.get('status', '就餐')
            except:
                pass

        all_dates.sort()

        header = ['姓名'] + all_dates + ['本月就餐次数']
        ws.append(header)

        for name in sorted(all_students.keys()):
            date_statuses = all_students[name]
            eat_count = sum(1 for d in all_dates if date_statuses.get(d) == '就餐')

            row_data = [name] + [date_statuses.get(d, '') for d in all_dates] + [eat_count]
            ws.append(row_data)

    else:
        # ---------- 表现：按学生分组，每条记录一行，姓名和总分合并居中 ----------
        from openpyxl.styles import Alignment

        # 收集每个学生的表现记录
        perf_data = {}  # name -> [(date, status, reason), ...]
        for row in rows:
            try:
                records_list = json.loads(row['data'])
                for r in records_list:
                    name = r.get('name', '').strip()
                    if not name or name not in valid_students:
                        continue
                    status = r.get('status', '')
                    if status == '正常':
                        continue
                    reason = r.get('reason', '')
                    if name not in perf_data:
                        perf_data[name] = []
                    perf_data[name].append((row['date'], status, reason))
            except:
                pass

        # 表头
        ws.append(['姓名', '时间', '表现', '原因', '总分'])
        ws.append([])  # 空行分隔

        # 按姓名排序输出
        for name in sorted(perf_data.keys()):
            records = perf_data[name]
            records.sort(key=lambda x: x[0])  # 按日期排序
            base_score = 100
            praise_count = sum(1 for r in records if r[1] == '表扬')
            critic_count = sum(1 for r in records if r[1] == '批评')
            total_score = base_score + praise_count - critic_count

            # 第一行：姓名和总分合并，跨所有记录行
            row_count = len(records)
            start_row = ws.max_row + 1

            # 写入该学生第一行（只有姓名和总分，内容在后续行）
            for i, (date, status, reason) in enumerate(records):
                row_num = start_row + i
                if i == 0:
                    # 姓名（跨row_count行）
                    ws.cell(row=row_num, column=1, value=name)
                else:
                    ws.cell(row=row_num, column=1, value='')
                ws.cell(row=row_num, column=2, value=date)
                ws.cell(row=row_num, column=3, value=status)
                ws.cell(row=row_num, column=4, value=reason)
                ws.cell(row=row_num, column=5, value='')

            # 合并姓名列（跨所有记录行）
            ws.merge_cells(start_row=start_row, start_column=1, end_row=start_row + row_count - 1, end_column=1)
            ws.cell(row=start_row, column=1).alignment = Alignment(horizontal='center', vertical='center')

            # 合并总分列（跨所有记录行）
            ws.merge_cells(start_row=start_row, start_column=5, end_row=start_row + row_count - 1, end_column=5)
            ws.cell(row=start_row, column=5, value=total_score)
            ws.cell(row=start_row, column=5).alignment = Alignment(horizontal='center', vertical='center')

            ws.append([])  # 空行分隔不同学生

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    filename = f'{module_name}记录_{date_from}_{date_to}.xlsx'
    return send_file(buf, mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True, download_name=filename)

@app.route('/api/export/view', methods=['GET'])
def export_view():
    """在浏览器中直接显示数据（不下载）"""
    module = request.args.get('module', '')
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    if not date_from or not date_to:
        return '请提供日期区间', 400

    conn = get_db()
    sql = 'SELECT * FROM records WHERE 1=1'
    params = []
    if module:
        sql += ' AND module = ?'
        params.append(module)
    sql += ' AND date >= ? AND date <= ?'
    params.extend([date_from, date_to])
    sql += ' ORDER BY date ASC'
    rows = conn.execute(sql, params).fetchall()
    conn.close()

    module_names = {'attendance': '考勤', 'dining': '就餐', 'performance': '表现'}
    module_name = module_names.get(module, module)

    # 构建HTML表格
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>{module_name}记录</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;padding:12px}}
h2{{text-align:center;color:#333;margin-bottom:12px;font-size:18px}}
.info{{text-align:center;color:#666;font-size:13px;margin-bottom:12px}}
table{{width:100%;border-collapse:collapse;background:#fff;border-radius:8px;overflow:hidden;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
th,td{{padding:8px 10px;border-bottom:1px solid #eee;text-align:center;font-size:13px}}
th{{background:#fafafa;font-weight:600;color:#333}}
td:first-child,td:nth-child(2){{text-align:center}}
td:nth-child(3){{text-align:left}}
tr:last-child td{{border-bottom:none}}
.btn{{display:inline-block;padding:8px 16px;background:#1890ff;color:#fff;border-radius:8px;text-decoration:none;font-size:13px;margin:12px auto;text-align:center}}
</style>
</head>
<body>
<h2>📊 {module_name}记录</h2>
<p class="info">{date_from} 至 {date_to}</p>
<table>
<thead>
<tr>
'''
    if module == 'attendance':
        html += '<th>开始日期</th><th>结束日期</th><th>姓名</th><th>考勤情况</th>'
    elif module == 'dining':
        html += '<th>开始日期</th><th>结束日期</th><th>姓名</th><th>就餐情况</th>'
    elif module == 'performance':
        html += '<th>日期</th><th>姓名</th><th>表现</th><th>原因</th>'
    html += '</tr></thead><tbody>'

    if module == 'attendance':
        # 按学生聚合：统计整个区间的病假/事假/迟到（保留所有日期的数据）
        student_data = {}  # name -> {sick:0, personal:0, late:0}
        for row in rows:
            # 移除 is_workday 过滤：考勤记录可能出现在任何日期
            try:
                records_list = json.loads(row['data'])
                for r in records_list:
                    name = r.get('name', '')
                    status = r.get('status', '')
                    if not name:
                        continue
                    if name not in student_data:
                        student_data[name] = {'sick': 0, 'personal': 0, 'late': 0}
                    if status in ('病假0.5日', '病假1日'):
                        student_data[name]['sick'] += 0.5 if '0.5' in status else 1
                    elif status in ('事假0.5日', '事假1日'):
                        student_data[name]['personal'] += 0.5 if '0.5' in status else 1
                    elif status == '迟到':
                        student_data[name]['late'] += 1
            except:
                pass
        for name in sorted(student_data.keys()):
            d = student_data[name]
            parts = []
            if d['sick'] == 0 and d['personal'] == 0 and d['late'] == 0:
                status_text = '全勤'
            else:
                if d['sick'] > 0:
                    parts.append(f"病假{d['sick']}天")
                if d['personal'] > 0:
                    parts.append(f"事假{d['personal']}天")
                if d['late'] > 0:
                    parts.append(f"迟到{d['late']}次")
                status_text = ' '.join(parts)
            html += f"<tr><td>{date_from}</td><td>{date_to}</td><td>{name}</td><td style='text-align:left'>{status_text}</td></tr>"

    elif module == 'dining':
        # 按学生聚合：统计整个区间的就餐次数（保留所有日期的数据）
        student_data = {}  # name -> {total: 0, days: set()}
        for row in rows:
            # 移除 is_workday 过滤：就餐记录可能出现在任何日期
            try:
                records_list = json.loads(row['data'])
                for r in records_list:
                    name = r.get('name', '')
                    status = r.get('status', '')
                    if not name:
                        continue
                    if name not in student_data:
                        student_data[name] = {'total': 0, 'days': set()}
                    if status == '就餐':
                        student_data[name]['total'] += 1
                        student_data[name]['days'].add(row['date'])
            except:
                pass
        for name in sorted(student_data.keys()):
            count = student_data[name]['total']
            html += f"<tr><td>{date_from}</td><td>{date_to}</td><td>{name}</td><td style='text-align:left'>就餐{count}次</td></tr>"

    elif module == 'performance':
        for row in rows:
            try:
                records_list = json.loads(row['data'])
                for r in records_list:
                    html += f"<tr><td>{row['date']}</td><td>{r.get('name','')}</td><td>{r.get('status','')}</td><td>{r.get('reason','')}</td></tr>"
            except:
                pass

    html += '</tbody></table>'
    html += '<div style="text-align:center"><a href="javascript:window.close()" class="btn">← 返回上一页</a></div>'
    html += '</body></html>'

    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}

@app.route('/api/export/term', methods=['GET'])
def export_term():
    """按学期导出数据"""
    term = request.args.get('term', '')
    if not term:
        return jsonify({'error': '请提供学期参数，如 term=2025-2026-2'}), 400
    parts = term.split('-')
    if len(parts) != 3:
        return jsonify({'error': '学期格式错误，如 2025-2026-2'}), 400
    year_start, year_end, term_num = int(parts[0]), int(parts[1]), parts[2]
    if term_num == '1':
        d1, d2 = f'{year_start}-09-01', f'{year_end}-01-20'
    else:
        d1, d2 = f'{year_end}-02-01', f'{year_end}-07-15'

    conn = get_db()
    records_data = conn.execute(
        'SELECT * FROM records WHERE date >= ? AND date <= ? ORDER BY date', (d1, d2)
    ).fetchall()
    students_data = conn.execute('SELECT * FROM students ORDER BY id').fetchall()
    conn.close()
    return jsonify({
        'term': term, 'students': [dict(r) for r in students_data],
        'records': [dict(r) for r in records_data]
    })

# ---------- 家长沟通 ----------
def _generate_parent_content(topic, student_name, parent_name):
    """根据沟通主题AI生成沟通内容摘要（约300字，占半页A4）"""
    t = topic.strip()

    templates = {

        '上课注意力不集中': f"""一、沟通主题：上课注意力不集中
本次与{parent_name}就{student_name}在课堂上注意力不集中、小动作较多的问题进行了深入探讨。

二、在校表现
{student_name}在课堂上容易分心，对感兴趣的科目能保持专注，但整体专注时长较短。老师和同学曾多次提醒其专心听讲，但效果有限。

三、成因分析
1. 孩子年龄特点，好奇心强，容易被新事物吸引
2. 课堂内容可能存在跟不上或已掌握的情况
3. 家庭环境中电子产品使用时间较长

四、解决建议
1. 在家设定固定学习时间，减少干扰，训练专注力
2. 每天与孩子交流"今天课堂上最有意思的事"，帮助其回顾课堂内容
3. 减少电子产品使用时间，建议每天不超过1小时
4. 尝试舒尔特训练法等专注力游戏
5. 与孩子约定进步目标，以正向激励为主

五、后续跟进
双方商定每两周沟通一次孩子的进步情况，共同帮助{student_name}提升专注力。""",

        '不写作业/拖欠作业': f"""一、沟通主题：不写作业/拖欠作业
本次与{parent_name}就{student_name}经常不完成或拖延作业的问题进行了深入交流。

二、在校表现
{student_name}在班级中属于聪明活泼的孩子，但作业完成情况不稳定，有时能独立完成，有时字迹潦草甚至不交作业。

三、成因分析
1. 学习习惯尚未养成，放学后缺乏有效监督
2. 作业有难度时缺乏迎难而上的意志力
3. 家庭环境中缺乏固定的学习氛围

四、解决建议
1. 每天放学后安排固定作业时间（建议18:00-20:00），家长陪伴但不代劳
2. 先完成作业再玩手机或看电视，形成正向激励
3. 遇到难题时，鼓励孩子先独立思考，实在不会的第二天问老师
4. 建立作业完成打卡表，连续完成一周给予适当奖励
5. 及时与班主任沟通，了解每日作业内容

五、后续跟进
{parent_name}表示会积极配合，共同帮助{student_name}养成按时完成作业的好习惯。""",

        '打架/与同学冲突': f"""一、沟通主题：打架/与同学冲突
本次与{parent_name}就{student_name}与同学发生冲突、甚至动手打架的问题进行了严肃而深入的交流。

二、事件经过
{student_name}与同学发生矛盾时，情绪容易激动，未能控制好自己的言行，导致肢体冲突，对双方都造成了不良影响。

三、成因分析
1. 情绪管理能力不足，遇到冲突时缺乏冷静思考
2. 语言表达和沟通能力有待提升
3. 可能存在模仿家庭成员处理问题的方式

四、教育建议
1. 与孩子深入交谈，了解事件经过，倾听孩子的感受，不急于批评
2. 明确告知打架是错误行为，并解释其可能产生的后果
3. 教给孩子正确的情绪表达方式和冲突解决方法
4. 引导{student_name}主动向被打同学道歉，培养责任感
5. 请家长在家也避免用武力解决问题，为孩子做好榜样

五、后续跟进
双方商定持续关注{student_name}的情绪变化和行为表现，争取彻底改正。""",

        '迟到/上学迟到': f"""一、沟通主题：迟到/上学迟到
本次与{parent_name}就{student_name}多次上学迟到的问题进行了交流。

二、在校表现
{student_name}经常踩着上课铃声进教室，甚至有时会迟到几分钟。虽然次数不多，但已引起班级老师和同学的注意。

三、成因分析
1. 作息时间不规律，晚上睡得晚，早上起不来
2. 早晨时间安排混乱，缺乏固定出门时间
3. 可能存在对学校或某个科目抵触的情绪

四、解决建议
1. 调整作息时间，确保每天睡眠不少于9小时
2. 制定"早晨时间表"，提前15分钟出门
3. 设立"准时奖励"，一周不迟到给予小奖励
4. 了解孩子是否有厌学情绪，关心其在校状态
5. 前一晚准备好第二天衣物书包，减少早晨忙乱

五、后续跟进
{parent_name}表示会重视孩子作息管理，配合学校共同解决迟到问题。""",

        '早恋/与异性交往过密': f"""一、沟通主题：早恋/与异性交往过密
本次与{parent_name}就{student_name}与异性同学交往过密，可能存在早恋倾向的问题进行了谨慎而真诚的交流。

二、观察情况
老师观察到{student_name}近期情绪波动较大，上课走神，与某位异性同学课间互动频繁，回家后手机聊天记录增多。

三、成因分析
1. 青春期生理心理发展，对异性产生好奇和好感是正常现象
2. 学业压力大，通过情感寄托获得心理慰藉
3. 对爱情和人际关系认知尚不成熟

四、教育建议
1. 请家长保持冷静，不要简单粗暴地禁止或批评，避免激化矛盾
2. 以朋友的方式与孩子谈心，了解其真实想法
3. 正面引导：告诉孩子中学阶段以学业为重，情感问题是人生必经阶段但要学会处理
4. 鼓励孩子扩大朋友圈，多参与集体活动，避免陷入一对一情感纠葛
5. 关注孩子情绪变化，避免因情感问题影响学业和心理健康

五、后续跟进
双方共同关注{student_name}的情感动态，保持顺畅沟通。""",

        '沉迷网络游戏': f"""一、沟通主题：沉迷网络游戏
本次与{parent_name}就{student_name}沉迷网络游戏，影响学习和作息的问题进行了深入探讨。

二、表现情况
{student_name}在家经常偷偷玩手机游戏，有时玩到深夜，影响第二天的学习状态。近期学习成绩有所下滑，精神状态不佳。

三、成因分析
1. 课余生活单一，缺乏其他兴趣爱好
2. 同学间存在游戏社交圈子，有从众心理
3. 家长监管力度不足，手机管理不严

四、解决建议
1. 与孩子坦诚交流，了解其游戏内容和交友圈，不完全否定游戏
2. 制定游戏时间规则：每周不超过2小时，周末可适当放宽
3. 将课余时间引导到体育运动、阅读、音乐等其他兴趣上
4. 请家长以身作则，在孩子面前减少玩手机的时间
5. 设置青少年模式或家长控制软件辅助管理

五、后续跟进
{parent_name}表示会重视并调整管理方式，双方约定定期交流孩子情况。""",

        '成绩明显下滑': f"""一、沟通主题：成绩明显下滑
本次与{parent_name}就{student_name}近期学习成绩明显下滑的问题进行了深入分析。

二、成绩现状
与期中考试相比，{student_name}本次测验成绩下降了约15分，主要失分在基础题和阅读理解部分。

三、成因分析
1. 学习方法不够科学，缺少错题整理和复习巩固环节
2. 近阶段课堂听讲效率有所下降
3. 可能存在偏科现象，对部分科目缺乏兴趣

四、改进建议
1. 建立错题本，将每次作业和测验中的错题整理归类
2. 每天复习当天所学内容，每周进行一次知识梳理
3. 了解各科目学习状态，找出薄弱环节重点突破
4. 设定阶段性小目标，每次进步一点点，及时给予肯定
5. 如有需要可报名课外辅导或与老师约定课后辅导时间

五、后续跟进
双方约定密切关注{student_name}的学习状态，每月沟通一次学习情况。""",

        '偏科/学科薄弱': f"""一、沟通主题：偏科/学科薄弱
本次与{parent_name}就{student_name}在部分科目上存在明显偏科的问题进行了交流。

二、偏科现状
{student_name}在文科（语文、英语）方面表现较好，但对理科（数学、物理）明显缺乏兴趣和信心，每次考试理科成绩明显低于文科。

三、成因分析
1. 基础薄弱，跟不上课堂节奏，导致恶性循环
2. 对理科存在畏难情绪，心理上排斥
3. 花费在文科上的时间多，理科练习不足

四、解决建议
1. 从基础抓起，每天坚持做10-15分钟理科基础题
2. 发现理科中的趣味点，激发学习兴趣
3. 安排理科辅导时间，查漏补缺，夯实基础
4. 多给予正向鼓励，帮助建立理科学习的信心
5. 各科目均衡发展，避免过度偏科影响升学

五、后续跟进
{parent_name}表示会配合老师，帮助{student_name}逐步提升理科成绩。""",

        '亲子沟通/亲子关系紧张': f"""一、沟通主题：亲子沟通/亲子关系紧张
本次与{parent_name}就{student_name}与家长之间亲子沟通不畅的问题进行了交流。

二、表现情况
{student_name}在家与父母交流较少，有时家长询问学校情况时，孩子表现出不耐烦或沉默。亲子之间时有争执，关系略显紧张。

三、成因分析
1. 青春期孩子独立意识增强，不愿意事事向家长汇报
2. 家长沟通方式较为说教，孩子产生逆反心理
3. 家长对孩子期望值较高，孩子感到压力大

四、改善建议
1. 改变沟通方式：少用"质问"语气，多用"分享"方式和孩子聊天
2. 每天留出15-30分钟"专属时光"，放下手机，全身心陪伴
3. 尊重孩子的隐私和个人空间，给予适当信任
4. 避免在吃饭或睡前批评孩子，创造温馨家庭氛围
5. 定期召开家庭会议，让孩子参与家庭事务讨论

五、后续跟进
{parent_name}意识到需要改变沟通方式，愿意尝试新的亲子交流方法。""",

        '心理健康/情绪问题': f"""一、沟通主题：心理健康/情绪问题
本次与{parent_name}就{student_name}近期情绪低落、压力过大等心理健康问题进行了关切交流。

二、表现情况
{student_name}近期情绪波动明显，容易焦虑和情绪低落，学习效率下降，与同学交往减少，有时会说"没意思"等消极话语。

三、成因分析
1. 学业压力过大，尤其是期中期末考试前后
2. 人际关系出现问题，如与同学发生矛盾、被孤立等
3. 家庭环境变化（父母工作忙、家庭矛盾等）也可能是诱因

四、支持建议
1. 请家长多关注孩子的情绪变化，避免给过大学业压力
2. 每天与孩子进行简短的情感交流，了解其内心想法
3. 鼓励孩子发展兴趣爱好，适度运动帮助缓解压力
4. 如果情况持续或加重，建议寻求专业心理咨询师的帮助
5. 班主任会在学校持续关注{student_name}的情绪状态

五、后续跟进
双方约定密切关注孩子心理状况，保持密切沟通，如有必要及时寻求专业帮助。""",

        '同学矛盾/人际交往问题': f"""一、沟通主题：同学矛盾/人际交往问题
本次与{parent_name}就{student_name}与同学之间发生矛盾和冲突的问题进行了交流。

二、事件经过
{student_name}与班级同学因小事发生争执，双方互不相让，导致矛盾升级，对双方情绪都产生了不良影响。

三、处理情况
老师已分别与双方学生谈话，了解事情经过，进行了调解教育，并安排双方当面道歉，握手言和。

四、教育建议
1. 引导{student_name}理解"退一步海阔天空"的道理
2. 教给孩子正确处理人际矛盾的方法：先冷静，再沟通，必要时请老师帮忙
3. 鼓励{student_name}主动修复与同学的关系
4. 请家长也引导孩子学会宽容和理解他人
5. 关注孩子情绪，避免因人际关系问题影响学习和生活

五、后续跟进
{student_name}已认识到自己的问题，愿意主动改善与同学的关系。""",

    }

    # 预定义主题（用户从下拉选择时）
    predefined_topics = {
        '上课注意力不集中': templates['上课注意力不集中'],
        '不写作业/拖欠作业': templates['不写作业/拖欠作业'],
        '打架/与同学冲突': templates['打架/与同学冲突'],
        '迟到/上学迟到': templates['迟到/上学迟到'],
        '早恋/与异性交往过密': templates['早恋/与异性交往过密'],
        '沉迷网络游戏': templates['沉迷网络游戏'],
        '成绩明显下滑': templates['成绩明显下滑'],
        '偏科/学科薄弱': templates['偏科/学科薄弱'],
        '亲子沟通/亲子关系紧张': templates['亲子沟通/亲子关系紧张'],
        '心理健康/情绪问题': templates['心理健康/情绪问题'],
        '同学矛盾/人际交往问题': templates['同学矛盾/人际交往问题'],
    }

    # 1. 精确匹配预定义主题
    if t in predefined_topics:
        return predefined_topics[t]

    # 2. 按优先级关键词匹配
    keyword_rules = [
        (['打架', '动手打人', '打人'], '打架/与同学冲突'),
        (['沉迷游戏', '沉迷网络游戏', '游戏成瘾', '网瘾'], '沉迷网络游戏'),
        (['早恋', '恋爱', '喜欢异性', '交往过密'], '早恋/与异性交往过密'),
        (['迟到', '上学迟到', '晚到'], '迟到/上学迟到'),
        (['偏科', '学科薄弱', '理科弱', '文科弱'], '偏科/学科薄弱'),
        (['亲子', '亲子关系', '亲子沟通'], '亲子沟通/亲子关系紧张'),
        (['心理健康', '情绪问题', '心理问题', '抑郁', '焦虑'], '心理健康/情绪问题'),
        (['同学矛盾', '人际问题', '人际交往', '和同学'], '同学矛盾/人际交往问题'),
        (['不写作业', '不写', '拖欠作业', '作业不完成', '作业拖欠'], '不写作业/拖欠作业'),
        (['上课', '注意力', '走神', '发呆', '不认真听讲', '上课不认真', '听课'], '上课注意力不集中'),
        (['成绩下滑', '成绩下降', '退步', '考差了'], '成绩明显下滑'),
    ]

    for keywords, template_key in keyword_rules:
        for kw in keywords:
            if kw in t:
                return templates[template_key]

    # 3. 学习相关（厌学、不喜欢学习等）
    if any(kw in t for kw in ['学习', '厌学', '不爱学习', '不喜欢学习', '学不进去', '学业']):
        return f"""一、沟通主题：{topic}
本次与{parent_name}就{student_name}学习兴趣不高、不喜欢学习的问题进行了深入交流。

二、在校表现
{student_name}在学习上表现出明显的抵触情绪，对某些科目提不起兴趣，课堂上参与度不高，作业完成质量不稳定。

三、原因分析
1. 学习内容可能存在较大困难，跟不上进度导致挫败感
2. 对某些科目缺乏兴趣，觉得枯燥无味
3. 学习方法不当，效率低下，努力却看不到成果
4. 缺乏明确的学习目标和动力

四、改进建议
1. 与孩子耐心沟通，了解其具体困难点和真实想法
2. 设定阶段性的学习目标，每达成一个小目标给予鼓励
3. 探索适合孩子的学习方法，如思维导图、费曼学习法等
4. 尝试将学习内容与实际生活联系，激发学习兴趣
5. 避免简单粗暴地批评，多用正向激励重塑学习信心
6. 必要时可寻求课外辅导，查漏补缺，夯实基础

五、后续跟进
双方商定每两周沟通一次学习情况，共同帮助{student_name}重新建立学习兴趣和信心。"""

    # 4. 兜底：通用自适应模板
    return f"""一、沟通主题：{topic}
本次与{parent_name}就{student_name}在「{topic}」方面的情况进行了深入交流。

二、现状分析
在与{parent_name}的沟通中，详细了解了{student_name}在「{topic}」这一方面的具体表现和困扰。通过交流，共同分析了问题的成因。

三、问题探讨
围绕「{topic}」这一主题，双方进行了认真探讨。{parent_name}反映了孩子在家中的相关表现，老师分享了孩子在校期间的情况，双方对问题的认识更加清晰。

四、解决建议
1. 请家长密切关注{student_name}在「{topic}」方面的变化，及时给予关注和支持
2. 与孩子进行平等、耐心的沟通，了解真实想法，避免简单批评
3. 家校保持密切联系，形成教育合力，共同帮助孩子克服困难
4. 设定合理的改进目标，以正向激励为主，逐步改善
5. 如问题持续或加重，建议寻求专业指导

五、后续跟进
双方商定定期回顾{student_name}的改变情况，争取在家校的共同努力下，让孩子在「{topic}」方面取得进步，健康成长。"""


@app.route('/api/parent/save', methods=['POST'])
def save_parent_record():
    """保存家长沟通记录"""
    data = request.json
    date = data.get('date', '')
    student_name = data.get('student_name', '').strip()
    parent_name = data.get('parent_name', '').strip()
    topic = data.get('topic', '').strip()

    if not date or not student_name or not parent_name or not topic:
        return jsonify({'error': '请填写完整信息'}), 400

    # 生成AI内容
    content = _generate_parent_content(topic, student_name, parent_name)

    record_id = str(uuid.uuid4())
    conn = get_db()
    conn.execute('''CREATE TABLE IF NOT EXISTS parent_records (
        id TEXT PRIMARY KEY, date TEXT, student_name TEXT, parent_name TEXT, topic TEXT, content TEXT, created TEXT
    )''')
    conn.execute('''INSERT INTO parent_records (id, date, student_name, parent_name, topic, content, created) VALUES (?, ?, ?, ?, ?, ?, ?)''',
                (record_id, date, student_name, parent_name, topic, content, datetime.now().isoformat()))
    conn.commit()
    conn.close()
    return jsonify({'id': record_id, 'ok': True})

@app.route('/api/parent/export/view', methods=['GET'])
def export_parent_view():
    """显示家长沟通记录导出页面"""
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    conn = get_db()
    if date_from and date_to:
        rows = conn.execute(
            'SELECT * FROM parent_records WHERE date>=? AND date<=? ORDER BY date ASC',
            (date_from, date_to)
        ).fetchall()
    else:
        rows = conn.execute('SELECT * FROM parent_records ORDER BY date ASC').fetchall()
    conn.close()

    records_html = ''
    for r in rows:
        records_html += f'''<div style="background:#fafafa;border:1px solid #eee;border-radius:10px;padding:16px;margin-bottom:16px">
          <div style="display:flex;justify-content:space-between;align-items:center;margin-bottom:8px">
            <span style="font-size:13px;color:#666">{r['date']}</span>
          </div>
          <div style="font-size:14px;margin-bottom:4px"><b>学生：</b>{r['student_name']}　<b>家长：</b>{r['parent_name']}</div>
          <div style="font-size:14px;margin-bottom:8px"><b>主题：</b>{r['topic']}</div>
          <div style="font-size:13px;color:#555;line-height:1.8;white-space:pre-wrap;background:#f5f5f5;padding:10px;border-radius:6px">{r['content']}</div>
        </div>'''

    total_count = len(rows)
    date_range = f'{date_from} 至 {date_to}' if (date_from and date_to) else '全部记录'

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>家长沟通记录预览</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;padding:16px}}
h2{{text-align:center;font-size:18px;color:#333;margin-bottom:8px}}
.info{{text-align:center;font-size:13px;color:#666;margin-bottom:20px}}
.card{{background:#fff;border-radius:12px;padding:16px;max-width:680px;margin:0 auto;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.btn{{display:block;padding:14px;background:#1890ff;color:#fff;border-radius:10px;text-decoration:none;text-align:center;font-size:15px;margin-top:16px}}
.back{{display:block;text-align:center;padding:12px;margin-top:12px;color:#1890ff;text-decoration:none;font-size:14px}}
</style>
</head>
<body>
<h2>💬 家长沟通记录</h2>
<p class="info">{date_range}　共 {total_count} 条记录</p>
<div class="card">
  {records_html if records_html else '<div style="text-align:center;color:#999;padding:32px">该时间范围内暂无记录</div>'}
  {'<a href="/api/parent/export/docx?date_from=' + date_from + '&date_to=' + date_to + '" class="btn" style="background:#52c41a">📄 导出Word文档</a>' if total_count > 0 else ''}
</div>
<a href="javascript:history.back()" class="back">← 返回上一页</a>
</body>
</html>'''
    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}

@app.route('/api/parent/export/docx')
def export_parent_docx():
    """导出家长沟通Word文档"""
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    conn = get_db()
    if date_from and date_to:
        rows = conn.execute(
            'SELECT * FROM parent_records WHERE date>=? AND date<=? ORDER BY date ASC',
            (date_from, date_to)
        ).fetchall()
    else:
        rows = conn.execute('SELECT * FROM parent_records ORDER BY date ASC').fetchall()
    conn.close()

    doc = Document()
    # 标题
    title_p = doc.add_paragraph()
    title_run = title_p.add_run('家长沟通记录')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('')

    for i, r in enumerate(rows):
        # 每条记录标题
        p_title = doc.add_paragraph()
        p_title.add_run(f'【记录{i+1}】').bold = True
        p_title.runs[0].font.size = Pt(13)
        p_title.paragraph_format.space_after = Pt(6)

        # 5项基本信息
        info_lines = [
            ('日　　期', r['date']),
            ('学生姓名', r['student_name']),
            ('家长姓名', r['parent_name']),
            ('沟通主题', r['topic']),
        ]
        for label, value in info_lines:
            p = doc.add_paragraph()
            p.add_run(f'{label}：').bold = True
            p.add_run(value)
            p.runs[0].font.size = Pt(11)
            p.runs[1].font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)

        # 沟通内容
        p_content_title = doc.add_paragraph()
        p_content_title.add_run('沟通内容：').bold = True
        p_content_title.runs[0].font.size = Pt(11)
        p_content_title.paragraph_format.space_after = Pt(2)

        for line in r['content'].strip().split('\n'):
            line = line.strip()
            if not line:
                continue
            p = doc.add_paragraph()
            r_run = p.add_run(line)
            r_run.font.size = Pt(11)
            p.paragraph_format.space_after = Pt(2)
            p.paragraph_format.left_indent = Inches(0.2)

        # 分隔线（留白，约占半页）
        doc.add_paragraph('')
        doc.add_paragraph('')

    out_path = os.path.join(BASE_DIR, f'家长沟通记录_{date_from}_{date_to}.docx')
    doc.save(out_path)
    return send_file(out_path,
                     as_attachment=True,
                     download_name=f'家长沟通记录_{date_from}_{date_to}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# ---------- 照片采集 ----------
@app.route('/api/photo/activity/upload', methods=['POST'])
def upload_activity_photo():
    """上传活动照片"""
    file = request.files.get('photo')
    activity = request.form.get('activity', '').strip()
    date = request.form.get('date', '')
    if not file:
        return jsonify({'error': '没有文件'}), 400
    if not activity:
        return jsonify({'error': '请先输入活动名称'}), 400
    photo_id = str(uuid.uuid4())
    ext = os.path.splitext(file.filename or '')[1].lower()
    if ext not in ('.jpg', '.jpeg', '.png', '.gif', '.webp', '.heic', '.heif'):
        ext = '.jpg'
    filename = f'{photo_id}{ext}'
    save_path = os.path.join(UPLOAD_DIR, filename)
    file.save(save_path)

    # 保存到数据库
    conn = get_db()
    conn.execute('''INSERT INTO photos (id, filename, created) VALUES (?, ?, ?)''',
                (photo_id, filename, datetime.now().isoformat()))
    # 同时写入活动照片记录
    conn.execute('''CREATE TABLE IF NOT EXISTS activity_photos (
        id TEXT PRIMARY KEY, activity TEXT, date TEXT, filename TEXT, created TEXT
    )''')
    conn.execute('INSERT INTO activity_photos (id, activity, date, filename, created) VALUES (?, ?, ?, ?, ?)',
                (photo_id, activity, date, filename, datetime.now().isoformat()))
    conn.commit()
    conn.close()
    return jsonify({'id': photo_id, 'filename': filename, 'activity': activity})

@app.route('/api/photo/activity/list', methods=['GET'])
def list_activity_photos():
    """列出所有活动（按日期或全部）"""
    date_filter = request.args.get('date', '')
    conn = get_db()
    if date_filter:
        rows = conn.execute(
            'SELECT activity, date, COUNT(*) as cnt FROM activity_photos WHERE date=? GROUP BY activity, date ORDER BY date DESC',
            (date_filter,)
        ).fetchall()
    else:
        rows = conn.execute(
            'SELECT activity, date, COUNT(*) as cnt FROM activity_photos GROUP BY activity, date ORDER BY date DESC'
        ).fetchall()
    conn.close()
    activities = [{'name': r['activity'], 'date': r['date'], 'count': r['cnt']} for r in rows]
    return jsonify({'activities': activities})

@app.route('/api/photo/activity/export/view', methods=['GET'])
def export_activity_view():
    """显示导出页面（列出所有活动，点击生成Word）"""
    conn = get_db()
    rows = conn.execute(
        'SELECT activity, date, COUNT(*) as cnt FROM activity_photos GROUP BY activity, date ORDER BY date DESC'
    ).fetchall()
    conn.close()

    activities_html = ''
    for r in rows:
        name_enc = r['activity']
        date_str = r['date']
        cnt = r['cnt']
        activities_html += f'''<div style="display:flex;align-items:center;justify-content:space-between;padding:16px;background:#fafafa;border:1px solid #eee;border-radius:10px;margin-bottom:10px">
          <div>
            <div style="font-size:15px;font-weight:500">{name_enc}</div>
            <div style="font-size:13px;color:#666;margin-top:4px">{date_str} · {cnt}张照片</div>
          </div>
          <a href="/api/photo/activity/export/docx?name={name_enc}&date={date_str}" style="padding:10px 18px;background:#1890ff;color:#fff;border-radius:8px;text-decoration:none;font-size:14px">📤 导出Word</a>
        </div>'''

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>导出活动照片</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;padding:16px}}
h2{{text-align:center;font-size:18px;color:#333;margin-bottom:20px}}
.card{{background:#fff;border-radius:12px;padding:16px;max-width:640px;margin:0 auto;box-shadow:0 1px 3px rgba(0,0,0,.1)}}
.tip{{font-size:13px;color:#666;text-align:center;margin-bottom:16px;padding:12px;background:#f0f7ff;border-radius:8px}}
.back{{display:block;text-align:center;padding:12px;margin-top:16px;color:#1890ff;text-decoration:none;font-size:14px}}
</style>
</head>
<body>
<h2>📷 活动照片导出</h2>
<div class="card">
  <div class="tip">点击活动名称即可生成包含所有照片的Word文档</div>
  {activities_html if activities_html else '<div style="text-align:center;color:#999;padding:32px">暂无活动记录</div>'}
</div>
<a href="javascript:window.close()" class="back">← 返回</a>
</body>
</html>'''
    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}

@app.route('/api/photo/activity/export/docx')
def export_activity_docx():
    """生成活动照片Word文档"""
    name = request.args.get('name', '')
    date = request.args.get('date', '')
    if not name:
        return '请提供活动名称', 400

    conn = get_db()
    rows = conn.execute(
        'SELECT filename FROM activity_photos WHERE activity=? AND date=? ORDER BY created ASC',
        (name, date)
    ).fetchall()
    conn.close()

    doc = Document()
    # 标题
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(f'活动记录：{name}')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    # 基本信息
    info_p = doc.add_paragraph()
    info_run = info_p.add_run(f'活动日期：{date}')
    info_run.font.size = Pt(12)
    doc.add_paragraph(f'照片数量：{len(rows)}张')
    doc.add_paragraph('')

    # 插入照片（按时间顺序）
    for i, row in enumerate(rows):
        photo_path = os.path.join(UPLOAD_DIR, row['filename'])
        if os.path.exists(photo_path):
            try:
                # 添加图片说明
                p = doc.add_paragraph()
                r = p.add_run(f'第 {i+1} 张')
                r.bold = True
                r.font.size = Pt(10)
                r.font.color.rgb = RGBColor(128, 128, 128)
                doc.add_picture(photo_path, width=Inches(5.5))
            except Exception as e:
                doc.add_paragraph(f'（照片{i+1}：{row["filename"]}）')
        else:
            doc.add_paragraph(f'（照片{i+1}：文件不存在）')

    out_path = os.path.join(BASE_DIR, f'活动记录_{name}_{date}.docx')
    doc.save(out_path)
    return send_file(out_path,
                     as_attachment=True,
                     download_name=f'活动记录_{name}_{date}.docx',
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

# ---------- 课程安排 ----------
@app.route('/api/course/upload', methods=['POST'])
def upload_course_schedule():
    """解析上传的课程表Excel"""
    import io
    from openpyxl import load_workbook

    file = request.files.get('file')
    if not file:
        return jsonify({'error': '没有文件'}), 400

    try:
        wb = load_workbook(io.BytesIO(file.read()))
        ws = wb.active
        courses = []

        # 尝试自动识别列：星期、课节、科目、教师
        # 格式1：第1列=星期(int 1-7), 第2列=课节(int), 第3列=科目, 第4列=教师
        # 格式2：第1列=星期文字, 第2列=课节, 第3列=科目, 第4列=教师
        weekday_map = {'周一':0,'周二':1,'周三':2,'周四':3,'周五':4,'周六':5,'周日':6,
                       '星期一':0,'星期二':1,'星期三':2,'星期四':3,'星期五':4,'星期六':5,'星期日':6,
                       '1':0,'2':1,'3':2,'4':3,'5':4,'6':5,'7':6}

        for row in ws.iter_rows(min_row=2, values_only=True):
            if not row[0]:
                continue
            # 解析星期
            w_val = str(row[0]).strip()
            if w_val in weekday_map:
                weekday = weekday_map[w_val]
            elif w_val.isdigit() and 1 <= int(w_val) <= 7:
                weekday = int(w_val) - 1
            else:
                continue  # 无法识别的星期，跳过

            # 解析课节
            period_val = row[1]
            if period_val is None:
                continue
            if isinstance(period_val, (int, float)):
                period = int(period_val)
            elif str(period_val).isdigit():
                period = int(period_val)
            else:
                period_str = str(period_val).strip()
                # 尝试提取数字
                import re
                m = re.search(r'\d+', period_str)
                period = int(m.group()) if m else 0
                if period == 0:
                    continue

            # 解析科目
            subject = str(row[2]).strip() if len(row) > 2 and row[2] else ''
            if not subject:
                continue

            # 解析教师
            teacher = str(row[3]).strip() if len(row) > 3 and row[3] else '待定'

            courses.append({
                'weekday': weekday,
                'period': period,
                'subject': subject,
                'teacher': teacher
            })

        return jsonify({'courses': courses, 'total': len(courses)})

    except Exception as e:
        return jsonify({'error': '解析失败：' + str(e)}), 400

@app.route('/api/course/template', methods=['GET'])
def download_course_template():
    """下载课程表模板"""
    from openpyxl import Workbook
    from io import BytesIO

    wb = Workbook()
    ws = wb.active
    ws.title = '课程表'
    ws.append(['星期', '课节', '科目', '任课教师'])
    # 示例数据（周一到周五，每天6节课）
    weekdays_text = ['周一', '周二', '周三', '周四', '周五']
    for w_idx, w_text in enumerate(weekdays_text):
        for period in range(1, 7):
            ws.append([w_text, period, '语文', '张老师'])
    # 周六周日各3节课
    for w_text in ['周六', '周日']:
        for period in range(1, 4):
            ws.append([w_text, period, '自习', '班主任'])

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)
    return send_file(buf,
                     mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                     as_attachment=True,
                     download_name='课程表模板.xlsx')

# ---------- 学情反馈 ----------
def _generate_learning_feedback(student_name, date_from, date_to, student_gender):
    """根据学生在校数据生成学情反馈内容（约400字，分3段）"""
    conn = get_db()
    att_records = conn.execute(
        "SELECT date, data FROM records WHERE module='attendance' AND date>=? AND date<=?",
        (date_from, date_to)
    ).fetchall()
    din_records = conn.execute(
        "SELECT date, data FROM records WHERE module='dining' AND date>=? AND date<=?",
        (date_from, date_to)
    ).fetchall()
    perf_records = conn.execute(
        "SELECT date, data FROM records WHERE module='performance' AND date>=? AND date<=?",
        (date_from, date_to)
    ).fetchall()
    parent_records = conn.execute(
        "SELECT date, topic, content FROM parent_records WHERE date>=? AND date<=? AND student_name=? ORDER BY date",
        (date_from, date_to, student_name)
    ).fetchall()
    conn.close()

    sick_days = personal_days = late_count = normal_count = 0
    for row in att_records:
        try:
            for r in json.loads(row['data']):
                if r.get('name') != student_name:
                    continue
                status = r.get('status', '正常')
                if status == '正常':
                    normal_count += 1
                elif status in ('病假0.5日', '病假1日'):
                    sick_days += 0.5 if '0.5' in status else 1
                elif status in ('事假0.5日', '事假1日'):
                    personal_days += 0.5 if '0.5' in status else 1
                elif status == '迟到':
                    late_count += 1
        except:
            pass

    eat_days = not_eat_days = out_school_days = 0
    for row in din_records:
        try:
            for r in json.loads(row['data']):
                if r.get('name') != student_name:
                    continue
                status = r.get('status', '就餐')
                if status == '就餐':
                    eat_days += 1
                elif status == '未就餐':
                    not_eat_days += 1
                elif status == '不在校就餐':
                    out_school_days += 1
        except:
            pass

    praise_records = []
    critic_records = []
    for row in perf_records:
        try:
            for r in json.loads(row['data']):
                if r.get('name') != student_name:
                    continue
                status = r.get('status', '')
                reason = r.get('reason', '')
                if status == '表扬':
                    praise_records.append((row['date'], reason))
                elif status == '批评':
                    critic_records.append((row['date'], reason))
        except:
            pass

    para1_parts = []
    total_days = normal_count + sick_days + personal_days + late_count
    if total_days > 0:
        att_summary = []
        if sick_days > 0:
            att_summary.append(f'病假{sick_days}日')
        if personal_days > 0:
            att_summary.append(f'事假{personal_days}日')
        if late_count > 0:
            att_summary.append(f'迟到{late_count}次')
        if att_summary:
            para1_parts.append(f'在{date_from}至{date_to}期间，该生共{total_days}天有考勤记录。' + '、'.join(att_summary) + '。')
        else:
            para1_parts.append(f'在{date_from}至{date_to}期间，该生出勤正常，无病假、事假及迟到记录。')
    else:
        para1_parts.append(f'在{date_from}至{date_to}期间，暂无考勤记录。')

    din_summary = []
    if eat_days > 0:
        din_summary.append(f'在校就餐{eat_days}次')
    if not_eat_days > 0:
        din_summary.append(f'未就餐{not_eat_days}次')
    if out_school_days > 0:
        din_summary.append(f'不在校就餐{out_school_days}次')
    if din_summary:
        para1_parts.append('就餐情况：' + '、'.join(din_summary) + '。')
    else:
        para1_parts.append('就餐情况：期间无就餐记录。')

    paragraph1 = ' '.join(para1_parts) if para1_parts else f'在{date_from}至{date_to}期间，暂无考勤和就餐记录。'

    para2_parts = []
    if praise_records:
        praise_reasons = [r[1] for r in praise_records if r[1]] or []
        praise_str = f'共获得{len(praise_records)}次表扬'
        if praise_reasons:
            praise_str += f'（' + '、'.join(set(praise_reasons[:3])) + '）'
        para2_parts.append(f'课堂表现方面，该生{len(praise_records)}次受到表扬，{praise_str}。')
    else:
        para2_parts.append('课堂表现方面，该生在近期表现记录中暂无表扬记录。')

    if critic_records:
        critic_reasons = [r[1] for r in critic_records if r[1]] or []
        critic_str = f'受到{len(critic_records)}次批评'
        if critic_reasons:
            critic_str += f'（' + '、'.join(set(critic_reasons[:3])) + '）'
        para2_parts.append(critic_str + '。')
    else:
        para2_parts.append('课堂表现方面，该生在近期无批评记录，表现良好。')

    if parent_records:
        topics = [r['topic'] for r in parent_records]
        para2_parts.append(f'家长沟通方面，本阶段与家长沟通{len(parent_records)}次，沟通主题涉及' + '、'.join(topics[:3]) + '。')
    else:
        para2_parts.append('家长沟通方面，本阶段暂无与家长的沟通记录。')

    paragraph2 = ' '.join(para2_parts)

    gender_text = '他' if student_gender == '男' else '她'
    att_score = 100 - sick_days - personal_days - late_count

    if att_score >= 95:
        att_advice = f'考勤方面表现良好，{gender_text}能保持稳定的出勤率，建议继续保持。'
    elif att_score >= 80:
        att_advice = f'考勤方面需要注意，{gender_text}有一定的缺勤情况，建议与家长配合关注。'
    else:
        att_advice = f'考勤方面需要重点关注，{gender_text}缺勤次数较多，建议与家长深入沟通了解原因。'

    if len(praise_records) > len(critic_records):
        perf_advice = f'表现方面，{gender_text}以正面表现为主，在班级中表现较优秀，建议继续鼓励并给予适当挑战。'
    elif len(critic_records) > 0 and len(praise_records) == 0:
        perf_advice = f'表现方面，{gender_text}近期批评次数较多，需要关注行为习惯的培养，建议与家长共同制定改进计划。'
    else:
        perf_advice = f'表现方面，{gender_text}整体表现一般，建议关注课堂参与度和作业完成情况。'

    if att_score >= 90 and len(critic_records) == 0 and len(praise_records) >= 1:
        overall = f'总体评价：{student_name}同学本阶段在校表现良好，各方面均无明显异常，建议继续保持当前状态，均衡发展。'
    elif att_score < 80 or len(critic_records) >= 3:
        overall = f'总体评价：{student_name}同学本阶段在出勤或纪律方面需要重点关注，建议班主任与家长保持密切沟通，共同帮助学生改善。'
    else:
        overall = f'总体评价：{student_name}同学本阶段整体表现基本稳定，建议在保持现有优点的同时，针对薄弱环节加以改进，家校配合促进学生健康成长。'

    paragraph3 = att_advice + ' ' + perf_advice + ' ' + overall

    return {'paragraph1': paragraph1, 'paragraph2': paragraph2, 'paragraph3': paragraph3}


@app.route('/api/feedback/generate', methods=['GET'])
def generate_feedback_docx():
    """生成学情反馈Word文档"""
    student_name = request.args.get('student_name', '').strip()
    date_from = request.args.get('date_from', '')
    date_to = request.args.get('date_to', '')

    if not student_name or not date_from or not date_to:
        return jsonify({'error': '请提供学生姓名、开始日期和结束日期'}), 400

    conn = get_db()
    student_row = conn.execute("SELECT name, gender FROM students WHERE name=?", (student_name,)).fetchone()
    if not student_row:
        conn.close()
        return jsonify({'error': f'学生"{student_name}"不在名单中，请先上传学生名单'}), 400
    student_gender = student_row['gender'] or '男'
    conn.close()

    content = _generate_learning_feedback(student_name, date_from, date_to, student_gender)

    doc = Document()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(f'{student_name}的学情反馈报告')
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_run.font.name = '方正小标宋简体'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    info_para = doc.add_paragraph()
    info_run = info_para.add_run(f'反馈周期：{date_from} 至 {date_to}')
    info_run.font.size = Pt(11)
    info_run.font.color.rgb = RGBColor(102, 102, 102)
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    p1_title = doc.add_paragraph()
    p1_run = p1_title.add_run('一、考勤与就餐情况')
    p1_run.bold = True
    p1_run.font.size = Pt(12)
    p1_run.font.color.rgb = RGBColor(0, 51, 102)
    p1_content = doc.add_paragraph()
    p1_c = p1_content.add_run(content['paragraph1'])
    p1_c.font.size = Pt(11)

    doc.add_paragraph('')

    p2_title = doc.add_paragraph()
    p2_run = p2_title.add_run('二、课堂表现与家校沟通')
    p2_run.bold = True
    p2_run.font.size = Pt(12)
    p2_run.font.color.rgb = RGBColor(0, 51, 102)
    p2_content = doc.add_paragraph()
    p2_c = p2_content.add_run(content['paragraph2'])
    p2_c.font.size = Pt(11)

    doc.add_paragraph('')

    p3_title = doc.add_paragraph()
    p3_run = p3_title.add_run('三、教育建议')
    p3_run.bold = True
    p3_run.font.size = Pt(12)
    p3_run.font.color.rgb = RGBColor(0, 51, 102)
    p3_content = doc.add_paragraph()
    p3_c = p3_content.add_run(content['paragraph3'])
    p3_c.font.size = Pt(11)

    doc.add_paragraph('')
    doc.add_paragraph('')

    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日")}')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(153, 153, 153)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    safe_name = student_name.replace(' ', '_').replace('/', '_')
    filename = f'学情反馈_{safe_name}_{date_from}_{date_to}.docx'
    out_path = os.path.join(BASE_DIR, filename)
    doc.save(out_path)

    return send_file(out_path,
                     as_attachment=True,
                     download_name=filename,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
def _generate_classlog_attendance(date_str, conn):
    att_row = conn.execute("SELECT data FROM records WHERE module='attendance' AND date=? ORDER BY created DESC LIMIT 1", (date_str,)).fetchone()
    class_total = conn.execute('SELECT COUNT(*) FROM students').fetchone()[0]
    sick_names = []
    personal_names = []
    late_names = []
    if att_row:
        records = json.loads(att_row[0])
        for r in records:
            s = r.get('status', '正常')
            if '病假' in s:
                sick_names.append(r['name'])
            elif '事假' in s:
                personal_names.append(r['name'])
            elif s == '迟到':
                late_names.append(r['name'])
    sick_count = len(sick_names)
    personal_count = len(personal_names)
    present = class_total - sick_count - personal_count
    rate = '良好' if (sick_count + personal_count) <= class_total * 0.05 else '基本正常，有少量缺勤需关注'
    sick_str = '、'.join(sick_names) if sick_names else '无'
    personal_str = '、'.join(personal_names) if personal_names else '无'
    late_count = len(late_names)
    detail_parts = []
    if sick_names: detail_parts.append(f"病假{sick_count}人（{sick_str}）")
    if personal_names: detail_parts.append(f"事假{personal_count}人（{personal_str}）")
    if late_names: detail_parts.append(f"迟到{late_count}人（{'、'.join(late_names)}）")
    detail_str = '；'.join(detail_parts) if detail_parts else '无异常情况，全员正常出勤'
    return f"一、考勤情况\n全班应到{class_total}人，实到{present}人。病假{sick_count}人（{sick_str}），事假{personal_count}人（{personal_str}），迟到{late_count}人。\n具体情况：{detail_str}。\n整体出勤状况{rate}。"

def _generate_classlog_dining(date_str, conn):
    din_row = conn.execute("SELECT data FROM records WHERE module='dining' AND date=? ORDER BY created DESC LIMIT 1", (date_str,)).fetchone()
    eat = present = 0
    not_eat_names = []
    if din_row:
        for r in json.loads(din_row[0]):
            s = r.get('status', '')
            if s == '就餐':
                eat += 1
                present += 1
            elif s == '未就餐':
                present += 1
                name = r.get('name', '')
                if name:
                    not_eat_names.append(name)
    rate_text = '98%以上，情况良好' if present == 0 or eat / present >= 0.98 else f'约{round(eat / max(present, 1) * 100)}%，需关注未就餐学生情况'
    not_eat_str = '无未就餐学生' if not not_eat_names else '未就餐学生：' + '、'.join(not_eat_names[:10])
    return f"二、就餐情况\n当日报到{present}人，实际就餐{eat}人，就餐率{rate_text}。\n{not_eat_str}。"

def _generate_classlog_performance(date_str, conn):
    perf_row = conn.execute("SELECT data FROM records WHERE module='performance' AND date=? ORDER BY created DESC LIMIT 1", (date_str,)).fetchone()
    praise_list = []
    critic_list = []
    if perf_row:
        for r in json.loads(perf_row[0]):
            s = r.get('status', '')
            reason = r.get('reason', '') or ''
            if s == '表扬':
                praise_list.append(r['name'] + ('（' + reason + '）' if reason else ''))
            elif s == '批评':
                critic_list.append(r['name'] + ('（' + reason + '）' if reason else ''))
    praise_count = len(praise_list)
    critic_count = len(critic_list)
    if praise_count == 0 and critic_count == 0:
        overall = '暂无奖惩记录，班级秩序正常'
    elif praise_count >= critic_count:
        overall = '表扬人次多于批评人次，班级氛围积极向上'
    else:
        overall = '批评人次较多，需关注行为习惯培养，加强正向引导'
    praise_str = '、'.join(praise_list[:8]) if praise_list else '无'
    critic_str = '、'.join(critic_list[:8]) if critic_list else '无'
    return f"三、学生表现情况\n当日班级表扬{praise_count}人次，批评{critic_count}人次。\n表扬学生：{praise_str}。\n批评学生：{critic_str}。\n整体表现{overall}。"

def _generate_classlog_course(date_str, conn):
    try:
        from datetime import datetime
        day_of_week = datetime.strptime(date_str, '%Y-%m-%d').weekday()
    except:
        day_of_week = 0
    try:
        courses = conn.execute("SELECT period, subject, teacher FROM courses WHERE weekday=?", (day_of_week,)).fetchall()
        if courses:
            course_list = [f"第{r['period']}节 {r['subject']}（{r['teacher']}）" for r in courses]
            return f"四、课程安排情况\n当日按课表上课，共{len(courses)}节课。具体安排：{'；'.join(course_list)}。各科教师按时到岗，教学秩序井然。"
    except:
        pass
    return f"四、课程安排情况\n当日课程按课表正常进行。（课程表数据暂未录入）"

def _generate_classlog_parent(date_str, conn):
    parent_rows = conn.execute("SELECT student_name, parent_name, topic, content FROM parent_records WHERE date=?", (date_str,)).fetchall()
    if not parent_rows:
        return "五、家长沟通情况\n当日无家长沟通记录。"
    records = []
    for r in parent_rows:
        topic = r['topic'] if r['topic'] else '常规沟通'
        content_short = r['content'][:50] if r['content'] else '详见记录'
        records.append(f"{r['student_name']}的{r['parent_name']}：{topic}（{content_short}）")
    return f"五、家长沟通情况\n当日与家长共沟通{len(parent_rows)}次：\n" + '；'.join(records)

def _generate_classlog_activity(date_str, conn):
    activity_rows = conn.execute("SELECT activity, date FROM activity_photos WHERE date=?", (date_str,)).fetchall()
    if not activity_rows:
        return "六、今日活动情况\n当日无活动记录，日常教学活动正常进行。"
    activities = list(set([r['activity'] for r in activity_rows]))
    return f"六、今日活动情况\n当日共开展{len(activities)}项活动：{'、'.join(activities)}。共拍摄{len(activity_rows)}张照片，活动记录完整。"

def _generate_classlog_work(date_str, conn):
    work_rows = conn.execute("SELECT data FROM records WHERE module='class_meeting' AND date=?", (date_str,)).fetchall()
    if work_rows:
        last = work_rows[-1]
        d = json.loads(last[0])
        content = d.get('content', '') or ''
        theme = d.get('theme', '') or ''
        if content:
            # 规范化序号为 1、2、3...
            import re
            lines = content.split('\n')
            numbered_lines = []
            count = 1
            for line in lines:
                stripped = line.strip()
                if stripped:
                    # 去掉常见序号前缀
                    stripped = re.sub(r'^(\d+[、\.．)）\]]\s*)|([一二三四五六七八九十]+[、\.]\s*)|([①-⑨]\s*)|([A-Za-z][\.、]\s*)', '', stripped)
                    if stripped:
                        numbered_lines.append(f'{count}、{stripped}')
                        count += 1
            numbered_content = '\n'.join(numbered_lines)
            return f"七、班务工作记录\n班会主题：{theme}。\n工作记录：\n{numbered_content}"
    return "七、班务工作记录\n当日班务工作正常开展，班级管理有序进行，无特殊事项记录。"


def _generate_classlog_content(date_str, items, conn):
    """根据日期和勾选项目生成班务日志内容"""
    section_generators = {
        'cl-attendance': _generate_classlog_attendance,
        'cl-dining': _generate_classlog_dining,
        'cl-performance': _generate_classlog_performance,
        'cl-course': _generate_classlog_course,
        'cl-parent': _generate_classlog_parent,
        'cl-activity': _generate_classlog_activity,
        'cl-work': _generate_classlog_work,
    }
    sections = []
    item_order = ['cl-attendance','cl-dining','cl-performance','cl-course','cl-parent','cl-activity','cl-work']
    for item in item_order:
        if item in items:
            gen = section_generators.get(item)
            if gen:
                sections.append(gen(date_str, conn))
    return '\n\n'.join(sections)

@app.route('/api/classlog/preview', methods=['GET'])
def classlog_preview():
    """显示班务日志预览页"""
    date_str = request.args.get('date', '').strip()
    items_str = request.args.get('items', '')

    if not date_str:
        return '<p style="text-align:center;padding:32px;color:#999">请提供日期参数</p>', 400

    items = items_str.split(',') if items_str else []
    item_labels = {
        'cl-attendance':'考勤情况','cl-dining':'就餐情况','cl-performance':'学生表现情况',
        'cl-course':'课程安排情况','cl-parent':'家长沟通情况','cl-activity':'今日活动情况','cl-work':'班务工作记录'
    }
    selected = [item_labels.get(i, i) for i in items]
    download_url = f'/api/classlog/download?date={date_str}&items={items_str}'

    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>班务日志预览</title>
<style>
*{{box-sizing:border-box;margin:0;padding:0}}
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;padding:16px}}
h2{{text-align:center;font-size:18px;color:#333;margin-bottom:8px}}
.info{{text-align:center;font-size:13px;color:#666;margin-bottom:20px}}
.card{{background:#fff;border-radius:12px;padding:16px;max-width:680px;margin:0 auto;box-shadow:0 1px 3px rgba(0,0,0,.1);text-align:center}}
.btn{{display:block;padding:14px;background:#722ed1;color:#fff;border-radius:10px;text-decoration:none;text-align:center;font-size:15px;margin-top:16px}}
.back{{display:block;text-align:center;padding:12px;margin-top:12px;color:#722ed1;text-decoration:none;font-size:14px}}
</style>
</head>
<body>
<h2>📝 班务日志</h2>
<p class="info">{date_str}　已选择 {len(selected)} 项内容</p>
<div class="card">
  <p style="font-size:14px;color:#555;line-height:1.8">生成内容：{' / '.join(selected)}</p>
  <a href="{download_url}" class="btn">📄 确认下载Word文档</a>
</div>
<a href="javascript:history.back()" class="back">← 返回上一页</a>
</body>
</html>'''
    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}


@app.route('/api/classlog/download', methods=['GET'])
def classlog_download():
    """下载班务日志Word文档"""
    date_str = request.args.get('date', '').strip()
    items_str = request.args.get('items', '')
    if not date_str:
        return jsonify({'error': '请提供日期'}), 400

    items = items_str.split(',') if items_str else []

    conn = get_db()
    content = _generate_classlog_content(date_str, items, conn)
    conn.close()

    doc = Document()
    title_para = doc.add_paragraph()
    title_run = title_para.add_run(date_str + '班务日志')
    title_run.bold = True
    title_run.font.size = Pt(18)
    title_run.font.name = '方正小标宋简体'
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')

    body_para = doc.add_paragraph()
    body_run = body_para.add_run(content)
    body_run.font.size = Pt(11)
    body_run.font.name = '仿宋'

    doc.add_paragraph('')
    footer_para = doc.add_paragraph()
    footer_run = footer_para.add_run(f'生成时间：{datetime.now().strftime("%Y年%m月%d日 %H:%M")}')
    footer_run.font.size = Pt(10)
    footer_run.font.color.rgb = RGBColor(153, 153, 153)
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    filename = date_str + '班务日志.docx'
    out_path = os.path.join(BASE_DIR, filename)
    doc.save(out_path)

    # 返回HTML页面，微信浏览器可直接在新窗口打开文件
    filename_encoded = urllib.parse.quote(filename)
    file_url = f'/api/classlog/file?filename={filename_encoded}'
    html = f'''<!DOCTYPE html>
<html lang="zh-CN">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width,initial-scale=1">
<title>班务日志下载</title>
<style>
body{{font-family:-apple-system,'PingFang SC','Microsoft YaHei',sans-serif;background:#f0f2f5;padding:16px;text-align:center}}
.card{{background:#fff;border-radius:12px;padding:32px;max-width:500px;margin:60px auto;box-shadow:0 2px 8px rgba(0,0,0,.1)}}
h3{{color:#333;font-size:18px;margin-bottom:12px}}
p{{color:#666;font-size:14px;margin-bottom:24px}}
.btn{{display:inline-block;padding:14px 32px;background:#722ed1;color:#fff;border-radius:10px;text-decoration:none;font-size:16px}}
.back{{display:block;margin-top:16px;color:#722ed1;text-decoration:none;font-size:14px}}
</style>
</head>
<body>
<div class="card">
  <h3>✅ 班务日志已生成</h3>
  <p>{date_str} 班务日志.docx</p>
  <a href="{file_url}" class="btn">📄 点击打开Word文档</a>
  <p style="font-size:12px;color:#999;margin-top:10px">如果文件没有自动打开，请点击上方按钮</p>
</div>
<a href="/api/classlog/preview?date={date_str}&items={items_str}" class="back">← 返回上一页</a>
</body>
</html>'''
    return html, 200, {'Content-Type': 'text/html; charset=utf-8'}


@app.route('/api/classlog/file', methods=['GET'])
def classlog_file():
    """提供docx文件下载/预览"""
    filename = request.args.get('filename', '')
    if not filename:
        return 'Not found', 404
    # 解码文件名
    filename_decoded = urllib.parse.unquote(filename)
    safe_date = filename_decoded.replace('班务日志.docx', '')
    out_path = os.path.join(BASE_DIR, filename_decoded)
    if not os.path.exists(out_path):
        return '文件不存在', 404
    return send_file(out_path,
                     as_attachment=False,
                     download_name=filename_decoded,
                     mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')


@app.route('/api/classlog/generate', methods=['GET'])
def generate_classlog():
    return '<p style="text-align:center;padding:32px;color:#999">请访问班务日志页面生成报告</p>', 200, {'Content-Type': 'text/html; charset=utf-8'}


if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=False)
